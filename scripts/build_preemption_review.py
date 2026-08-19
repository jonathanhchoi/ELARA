"""Build the Stage 02 preemption review as a polished Word document.

The Markdown input is a run-scoped, auditable build source. The DOCX output is
the versioned researcher-facing artifact. This script refuses to overwrite an
existing output so that ELARA's artifact versioning contract remains intact.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor


CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "CBD5E1"
WHITE = "FFFFFF"

REQUIRED_METADATA = (
    "title",
    "subtitle",
    "verdict",
    "recommended_disposition",
    "scoop_risk",
    "review_date",
    "recheck_date",
)
VERDICTS = {"preempted", "partially preempted", "open"}
SCOOP_RISKS = {"low", "moderate", "high"}
EXECUTIVE_SUMMARY_SECTION = "Executive summary"
ANNOTATED_MAP_SECTION = "Annotated map of closest work"
EXECUTIVE_SUMMARY_MAX_WORDS = 1200
EXECUTIVE_SUMMARY_MATCH_HEADING_PREFIX = "Closest match:"
EXECUTIVE_SUMMARY_REQUIRED_LABELS = (
    "Bottom line",
    "Intended contribution",
    "Closest threats",
    "Remaining contribution",
    "Recommended disposition",
    "Scoop risk and access gaps",
)
EXECUTIVE_SUMMARY_MATCH_REQUIRED_LABELS = (
    "What the work says",
    "Relevant scope and basis",
    "Preemption of the intended contribution",
    "What remains",
    "Evidence",
)
REQUIRED_SECTIONS = (
    EXECUTIVE_SUMMARY_SECTION,
    ANNOTATED_MAP_SECTION,
    "Verdict and flip conditions",
    "Positioning and lineage",
    "Scoop risk",
    "Search methods and saturation evidence",
    "Access limitations and manual search packet",
    "Review date and recheck",
)
PLACEHOLDER_RE = re.compile(r"\b(?:TODO-PREEMPTION|TBD|PLACEHOLDER)\b", re.IGNORECASE)
INLINE_RE = re.compile(
    r"(\[[^\]]+\]\(https?://[^)]+\)|\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)"
)
ORDERED_RE = re.compile(r"^\s*\d+[.)]\s+(.+)$")
BULLET_RE = re.compile(r"^\s*[-+*]\s+(.+)$")


@dataclass(frozen=True)
class Block:
    kind: str
    text: str = ""
    level: int = 0
    rows: tuple[tuple[str, ...], ...] = ()


def _decode_value(raw: str) -> str:
    value = raw.strip()
    if not value:
        return ""
    if value[0] in {'"', "'"}:
        if value[0] == '"':
            try:
                decoded = json.loads(value)
            except json.JSONDecodeError as exc:
                raise ValueError(f"invalid quoted metadata value: {value}") from exc
            if not isinstance(decoded, str):
                raise ValueError(f"metadata values must be strings: {value}")
            return decoded.strip()
        if len(value) < 2 or value[-1] != "'":
            raise ValueError(f"invalid quoted metadata value: {value}")
        return value[1:-1].replace("''", "'").strip()
    return value


def parse_source(text: str) -> tuple[dict[str, str], str]:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    if not lines or lines[0].strip() != "---":
        raise ValueError("source must begin with the preemption-review metadata block")
    try:
        end = next(index for index in range(1, len(lines)) if lines[index].strip() == "---")
    except StopIteration as exc:
        raise ValueError("metadata block has no closing --- line") from exc

    metadata: dict[str, str] = {}
    for line in lines[1:end]:
        if not line.strip() or line.lstrip().startswith("#"):
            continue
        key, separator, raw = line.partition(":")
        key = key.strip().lower()
        if not separator or not re.fullmatch(r"[a-z][a-z0-9_]*", key):
            raise ValueError(f"invalid metadata line: {line}")
        if key in metadata:
            raise ValueError(f"duplicate metadata key: {key}")
        metadata[key] = _decode_value(raw)

    missing = [key for key in REQUIRED_METADATA if not metadata.get(key)]
    if missing:
        raise ValueError("missing required metadata: " + ", ".join(missing))
    if PLACEHOLDER_RE.search(text):
        raise ValueError("source still contains an unresolved placeholder")
    verdict = metadata["verdict"].lower()
    if verdict not in VERDICTS:
        raise ValueError("verdict must be preempted, partially preempted, or open")
    metadata["verdict"] = verdict
    scoop_risk = metadata["scoop_risk"].lower()
    if scoop_risk not in SCOOP_RISKS:
        raise ValueError("scoop_risk must be low, moderate, or high")
    metadata["scoop_risk"] = scoop_risk
    try:
        review_date = date.fromisoformat(metadata["review_date"])
        recheck_date = date.fromisoformat(metadata["recheck_date"])
    except ValueError as exc:
        raise ValueError("review_date and recheck_date must use YYYY-MM-DD") from exc
    if recheck_date < review_date:
        raise ValueError("recheck_date may not precede review_date")

    body = "\n".join(lines[end + 1 :]).strip()
    if not body:
        raise ValueError("source body is empty")
    return metadata, body


def _split_table_row(line: str) -> tuple[str, ...]:
    stripped = line.strip().strip("|")
    return tuple(cell.strip() for cell in stripped.split("|"))


def _is_table_separator(line: str, columns: int) -> bool:
    cells = _split_table_row(line)
    return len(cells) == columns and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def _starts_explicit_block(lines: list[str], index: int) -> bool:
    stripped = lines[index].strip()
    if not stripped:
        return True
    if stripped.startswith(("```", ">")):
        return True
    if re.match(r"^#{1,4}\s+", stripped):
        return True
    if re.fullmatch(r"(?:-{3,}|\*{3,}|_{3,})", stripped):
        return True
    if ORDERED_RE.match(lines[index]) or BULLET_RE.match(lines[index]):
        return True
    if "|" in stripped and index + 1 < len(lines):
        header = _split_table_row(stripped)
        if len(header) >= 2 and _is_table_separator(lines[index + 1], len(header)):
            return True
    return False


def parse_blocks(body: str) -> list[Block]:
    lines = body.splitlines()
    blocks: list[Block] = []
    paragraph: list[str] = []

    def flush_paragraph() -> None:
        if paragraph:
            blocks.append(Block("paragraph", " ".join(item.strip() for item in paragraph)))
            paragraph.clear()

    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if not stripped:
            flush_paragraph()
            index += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            code: list[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code.append(lines[index])
                index += 1
            if index >= len(lines):
                raise ValueError("unclosed fenced code block")
            blocks.append(Block("code", "\n".join(code)))
            index += 1
            continue
        heading = re.match(r"^(#{1,4})\s+(.+?)\s*$", stripped)
        if heading:
            flush_paragraph()
            blocks.append(Block("heading", heading.group(2), len(heading.group(1))))
            index += 1
            continue
        if re.fullmatch(r"(?:-{3,}|\*{3,}|_{3,})", stripped):
            flush_paragraph()
            blocks.append(Block("rule"))
            index += 1
            continue
        if stripped.startswith(">"):
            flush_paragraph()
            quote: list[str] = [stripped[1:].strip()]
            index += 1
            while index < len(lines) and lines[index].strip():
                continuation = lines[index].strip()
                if continuation.startswith(">"):
                    quote.append(continuation[1:].strip())
                elif _starts_explicit_block(lines, index):
                    break
                else:
                    quote.append(continuation)
                index += 1
            blocks.append(Block("callout", " ".join(quote)))
            continue
        ordered = ORDERED_RE.match(line)
        if ordered:
            flush_paragraph()
            index += 1
            item = [ordered.group(1).strip()]
            while index < len(lines) and lines[index].strip():
                if _starts_explicit_block(lines, index):
                    break
                item.append(lines[index].strip())
                index += 1
            blocks.append(Block("ordered", " ".join(item)))
            continue
        bullet = BULLET_RE.match(line)
        if bullet:
            flush_paragraph()
            index += 1
            item = [bullet.group(1).strip()]
            while index < len(lines) and lines[index].strip():
                if _starts_explicit_block(lines, index):
                    break
                item.append(lines[index].strip())
                index += 1
            blocks.append(Block("bullet", " ".join(item)))
            continue
        if "|" in stripped and index + 1 < len(lines):
            header = _split_table_row(stripped)
            if len(header) >= 2 and _is_table_separator(lines[index + 1], len(header)):
                flush_paragraph()
                rows = [header]
                index += 2
                while index < len(lines) and "|" in lines[index] and lines[index].strip():
                    row = _split_table_row(lines[index])
                    if len(row) != len(header):
                        raise ValueError("Markdown table rows must have the same column count")
                    rows.append(row)
                    index += 1
                blocks.append(Block("table", rows=tuple(rows)))
                continue
        paragraph.append(line)
        index += 1
    flush_paragraph()
    return blocks


def validate_sections(blocks: list[Block]) -> None:
    level_two = [block.text for block in blocks if block.kind == "heading" and block.level == 2]
    cursor = 0
    for required in REQUIRED_SECTIONS:
        try:
            cursor = level_two.index(required, cursor) + 1
        except ValueError as exc:
            raise ValueError(f"missing or misordered required section: {required}") from exc

    summary_index = next(
        index
        for index, block in enumerate(blocks)
        if block.kind == "heading"
        and block.level == 2
        and block.text == EXECUTIVE_SUMMARY_SECTION
    )
    summary_end = next(
        index
        for index, block in enumerate(blocks[summary_index + 1 :], start=summary_index + 1)
        if block.kind == "heading" and block.level == 2
    )
    summary_blocks = blocks[summary_index + 1 : summary_end]
    summary_text = " ".join(
        block.text if block.kind != "table" else " ".join(" ".join(row) for row in block.rows)
        for block in summary_blocks
    )
    summary_words = len(re.findall(r"\b[\w][\w'-]*\b", summary_text))
    if not summary_words:
        raise ValueError("executive summary must contain substantive text")
    if summary_words > EXECUTIVE_SUMMARY_MAX_WORDS:
        raise ValueError(
            "executive summary exceeds "
            f"{EXECUTIVE_SUMMARY_MAX_WORDS} words ({summary_words} found)"
        )
    missing_labels = [
        label
        for label in EXECUTIVE_SUMMARY_REQUIRED_LABELS
        if not re.search(rf"\*\*{re.escape(label)}:\*\*", summary_text, re.IGNORECASE)
    ]
    if missing_labels:
        raise ValueError(
            "executive summary is missing required labeled content: "
            + ", ".join(missing_labels)
        )

    match_indices = [
        index
        for index, block in enumerate(summary_blocks)
        if block.kind == "heading"
        and block.level == 3
        and block.text.lower().startswith(EXECUTIVE_SUMMARY_MATCH_HEADING_PREFIX.lower())
    ]
    if not match_indices:
        raise ValueError(
            "executive summary must contain at least one "
            f"'{EXECUTIVE_SUMMARY_MATCH_HEADING_PREFIX} [full citation]' subsection"
        )
    for match_number, match_index in enumerate(match_indices, start=1):
        heading = summary_blocks[match_index].text
        citation = heading[len(EXECUTIVE_SUMMARY_MATCH_HEADING_PREFIX) :].strip()
        if not citation:
            raise ValueError(
                f"executive summary closest match {match_number} must include a full citation"
            )
        match_end = (
            match_indices[match_number]
            if match_number < len(match_indices)
            else len(summary_blocks)
        )
        match_text = " ".join(
            block.text
            if block.kind != "table"
            else " ".join(" ".join(row) for row in block.rows)
            for block in summary_blocks[match_index + 1 : match_end]
        )
        missing_match_labels = [
            label
            for label in EXECUTIVE_SUMMARY_MATCH_REQUIRED_LABELS
            if not re.search(rf"\*\*{re.escape(label)}:\*\*", match_text, re.IGNORECASE)
        ]
        if missing_match_labels:
            raise ValueError(
                f"executive summary closest match '{citation}' is missing required "
                "labeled content: "
                + ", ".join(missing_match_labels)
            )

    map_index = next(
        index
        for index, block in enumerate(blocks)
        if block.kind == "heading" and block.level == 2 and block.text == ANNOTATED_MAP_SECTION
    )
    verdict_index = next(
        index
        for index, block in enumerate(blocks[map_index + 1 :], start=map_index + 1)
        if block.kind == "heading" and block.level == 2
    )
    if not any(
        block.kind == "heading" and block.level == 3
        for block in blocks[map_index + 1 : verdict_index]
    ):
        raise ValueError("annotated map must contain at least one level-three work heading")


def _set_style_font(style, name: str, size: float, color: str, bold: bool = False) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)


def _set_keep_with_next(style) -> None:
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.keep_together = True


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    _set_style_font(normal, "Calibri", 11, "000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        _set_style_font(style, "Calibri", size, color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        _set_keep_with_next(style)

    title = styles.add_style("ELARA Report Title", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(title, "Calibri", 26, NAVY, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True

    subtitle = styles.add_style("ELARA Report Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(subtitle, "Calibri", 12.5, MUTED)
    subtitle.font.italic = True
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle.paragraph_format.line_spacing = 1.1
    subtitle.paragraph_format.keep_with_next = True

    callout = styles.add_style("ELARA Callout", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(callout, "Calibri", 10.5, NAVY)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.12)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.1

    table_text = styles.add_style("ELARA Table Text", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_font(table_text, "Calibri", 9.5, "000000")
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(2)
    table_text.paragraph_format.line_spacing = 1.05

def _format_run(
    run,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    *,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _add_hyperlink(paragraph, label: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    text = OxmlElement("w:t")
    text.text = label
    run.append(text)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        token = match.group(0)
        link = re.fullmatch(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
        if link:
            _add_hyperlink(paragraph, link.group(1), link.group(2))
        elif token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _format_run(run, "Consolas", 9.5, DARK_BLUE)
        elif token.startswith("*"):
            paragraph.add_run(token[1:-1]).italic = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _set_paragraph_box(paragraph, *, fill: str, border_color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border_color)
    borders.append(left)


def _next_numbering_id(numbering, tag: str, attribute: str) -> int:
    values = []
    for element in numbering.findall(qn(tag)):
        raw = element.get(qn(attribute))
        if raw is not None and raw.isdigit():
            values.append(int(raw))
    return max(values, default=0) + 1


def add_numbering_definition(doc: Document, *, ordered: bool) -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = _next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
    num_id = _next_numbering_id(numbering, "w:num", "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "decimal" if ordered else "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if ordered else "•")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.extend([tabs, indent])
    level.extend([start, num_format, level_text, suffix, justification, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend([level, number])
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167


def _table_properties(table, widths: list[int], *, indent: int = TABLE_INDENT_DXA) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    width = tbl_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.insert(0, width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    table_indent = tbl_pr.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        tbl_pr.append(table_indent)
    table_indent.set(qn("w:w"), str(indent))
    table_indent.set(qn("w:type"), "dxa")

    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", 80), ("left", 120), ("bottom", 80), ("right", 120)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), BORDER)
        borders.append(border)

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)

    for row in table.rows:
        for cell, value in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = tc_pr.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                tc_pr.append(tc_width)
            tc_width.set(qn("w:w"), str(value))
            tc_width.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_metadata_table(doc: Document, metadata: dict[str, str]) -> None:
    rows = [
        ("Verdict", metadata["verdict"].title()),
        ("Recommended disposition", metadata["recommended_disposition"]),
        ("Scoop risk", metadata["scoop_risk"].title()),
        ("Review date", metadata["review_date"]),
        ("Recommended recheck", metadata["recheck_date"]),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    widths = [2200, CONTENT_WIDTH_DXA - 2200]
    _table_properties(table, widths)
    for row, (label, value) in zip(table.rows, rows):
        label_cell, value_cell = row.cells
        _shade_cell(label_cell, LIGHT_BLUE)
        _shade_cell(value_cell, WHITE)
        label_p = label_cell.paragraphs[0]
        label_p.style = doc.styles["ELARA Table Text"]
        label_run = label_p.add_run(label)
        _format_run(label_run, "Calibri", 9.5, NAVY, bold=True)
        value_p = value_cell.paragraphs[0]
        value_p.style = doc.styles["ELARA Table Text"]
        add_inline(value_p, value)
        if label == "Verdict":
            for run in value_p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(NAVY)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def _content_widths(rows: tuple[tuple[str, ...], ...]) -> list[int]:
    maxima = [max(len(row[index]) for row in rows) for index in range(len(rows[0]))]
    weights = [max(12, min(45, value)) for value in maxima]
    raw = [round(CONTENT_WIDTH_DXA * weight / sum(weights)) for weight in weights]
    minimum = 900
    widths = [max(minimum, value) for value in raw]
    while sum(widths) > CONTENT_WIDTH_DXA:
        largest = max(range(len(widths)), key=lambda index: widths[index])
        if widths[largest] <= minimum:
            break
        widths[largest] -= min(60, sum(widths) - CONTENT_WIDTH_DXA)
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(doc: Document, rows: tuple[tuple[str, ...], ...]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    widths = _content_widths(rows)
    _table_properties(table, widths)
    _set_repeat_header(table.rows[0])
    for row_index, (row, values) in enumerate(zip(table.rows, rows)):
        for cell, value in zip(row.cells, values):
            _shade_cell(cell, LIGHT_BLUE if row_index == 0 else WHITE)
            paragraph = cell.paragraphs[0]
            paragraph.style = doc.styles["ELARA Table Text"]
            add_inline(paragraph, value)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(NAVY)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_rule(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(6)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), BORDER)
    borders.append(bottom)
    p_pr.append(borders)


def build_document(metadata: dict[str, str], blocks: list[Block]) -> Document:
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = metadata["title"]
    doc.core_properties.subject = "ELARA Stage 02 preemption review"
    doc.core_properties.author = "ELARA"
    doc.core_properties.keywords = "literature review, preemption review, empirical legal research"

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(4)
    kicker.paragraph_format.keep_with_next = True
    run = kicker.add_run("ELARA PREEMPTION REVIEW")
    _format_run(run, "Calibri", 9, BLUE, bold=True)
    title = doc.add_paragraph(metadata["title"], style="ELARA Report Title")
    title.paragraph_format.keep_with_next = True
    subtitle = doc.add_paragraph(metadata["subtitle"], style="ELARA Report Subtitle")
    subtitle.paragraph_format.keep_with_next = True
    add_metadata_table(doc, metadata)

    ordered_num: int | None = None
    bullet_num: int | None = None
    prior_kind = ""
    for block in blocks:
        if block.kind == "heading":
            if block.level == 1 and block.text.strip().lower() == metadata["title"].strip().lower():
                prior_kind = block.kind
                continue
            style_level = min(max(block.level - 1, 1), 3)
            paragraph = doc.add_paragraph(style=f"Heading {style_level}")
            if block.level == 2 and block.text == ANNOTATED_MAP_SECTION:
                paragraph.paragraph_format.page_break_before = True
            add_inline(paragraph, block.text)
        elif block.kind == "paragraph":
            paragraph = doc.add_paragraph()
            add_inline(paragraph, block.text)
        elif block.kind == "callout":
            paragraph = doc.add_paragraph(style="ELARA Callout")
            add_inline(paragraph, block.text)
            _set_paragraph_box(paragraph, fill=LIGHT_GRAY, border_color=BLUE)
        elif block.kind == "code":
            paragraph = doc.add_paragraph(style="ELARA Callout")
            run = paragraph.add_run(block.text)
            _format_run(run, "Consolas", 9, DARK_BLUE)
            _set_paragraph_box(paragraph, fill=LIGHT_GRAY, border_color=BORDER)
        elif block.kind in {"ordered", "bullet"}:
            if block.kind == "ordered" and prior_kind != "ordered":
                ordered_num = add_numbering_definition(doc, ordered=True)
            if block.kind == "bullet" and prior_kind != "bullet":
                bullet_num = add_numbering_definition(doc, ordered=False)
            paragraph = doc.add_paragraph()
            add_inline(paragraph, block.text)
            _apply_numbering(
                paragraph,
                ordered_num if block.kind == "ordered" else bullet_num,  # type: ignore[arg-type]
            )
        elif block.kind == "table":
            add_markdown_table(doc, block.rows)
        elif block.kind == "rule":
            add_rule(doc)
        prior_kind = block.kind
    return doc


def validate_output(path: Path) -> None:
    reopened = Document(path)
    if not reopened.tables:
        raise ValueError("generated document is missing its metadata table")
    headings = [
        paragraph.text.strip()
        for paragraph in reopened.paragraphs
        if paragraph.style is not None and paragraph.style.name == "Heading 1"
    ]
    missing = [section for section in REQUIRED_SECTIONS if section not in headings]
    if missing:
        raise ValueError("generated document is missing sections: " + ", ".join(missing))
    if any(PLACEHOLDER_RE.search(paragraph.text or "") for paragraph in reopened.paragraphs):
        raise ValueError("generated document contains an unresolved placeholder")


def build(source: Path, output: Path) -> None:
    if source.suffix.lower() != ".md":
        raise ValueError("source must be a .md file")
    if output.suffix.lower() != ".docx":
        raise ValueError("output must be a .docx file")
    if not source.is_file():
        raise ValueError(f"source does not exist: {source}")
    if output.exists():
        raise ValueError(f"refusing to overwrite existing artifact: {output}")
    metadata, body = parse_source(source.read_text(encoding="utf-8"))
    blocks = parse_blocks(body)
    validate_sections(blocks)
    output.parent.mkdir(parents=True, exist_ok=True)
    document = build_document(metadata, blocks)
    document.save(output)
    try:
        validate_output(output)
    except Exception:
        output.unlink(missing_ok=True)
        raise


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("source", type=Path, help="run-scoped Markdown build source")
    parser.add_argument("output", type=Path, help="new versioned .docx artifact")
    args = parser.parse_args()
    try:
        build(args.source.resolve(), args.output.resolve())
    except (OSError, ValueError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    print(f"Created {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
