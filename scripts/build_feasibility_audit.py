"""Build the Stage 03 feasibility audit as a question-led formatted report.

The Markdown input is the immutable, run-scoped build source. LaTeX is the
default output and compiles to the active PDF. DOCX remains available when the
researcher explicitly requests Word. The builder enforces the display questions,
rejects table-based reports, and refuses to overwrite an artifact.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from build_preemption_review import (
    BLUE,
    DARK_BLUE,
    LIGHT_GRAY,
    MUTED,
    NAVY,
    _apply_numbering,
    _format_run,
    _set_paragraph_box,
    add_inline,
    add_numbering_definition,
    configure_document,
    parse_blocks,
)
from latex_report import render_latex_report, validate_latex_output


@dataclass(frozen=True)
class GateSpec:
    internal_id: str
    question: str


GATES = (
    GateSpec("task-type", "Is the coding task one that LLMs are good at?"),
    GateSpec(
        "variable-verifiability",
        "Can a careful human verify each coding decision from the source?",
    ),
    GateSpec(
        "either-way-contribution",
        "Would this be an interesting contribution to the literature regardless of the direction of the results?",
    ),
    GateSpec("data-access", "Can we obtain and use the data the project needs?"),
    GateSpec(
        "base-rate-and-power",
        "Will there be enough usable data to answer the research question?",
    ),
    GateSpec(
        "time-and-resources",
        "Can the project be completed in a reasonable amount of time with the available resources?",
    ),
    GateSpec(
        "measurement-error",
        "Could coding errors change the answer, and can the analysis account for them?",
    ),
    GateSpec(
        "researcher-decision",
        "Does a legal, ethical, data-use, or spending issue require the researcher’s decision?",
    ),
)

DECISION_SUMMARY = "Decision summary"
BINDING_CONSTRAINT = "Binding constraint"
RECOMMENDATION = "Recommendation and what would change it"
LIMITATIONS = "Evidence gaps and limitations"
REQUIRED_SECTIONS = (
    DECISION_SUMMARY,
    *(gate.question for gate in GATES),
    BINDING_CONSTRAINT,
    RECOMMENDATION,
    LIMITATIONS,
)
REQUIRED_METADATA = (
    "title",
    "subtitle",
    "recommendation",
    "audit_date",
    "report_version",
)
RECOMMENDATIONS = {"go", "go with modifications", "no-go"}
GATE_DECISIONS = {"pass", "pass with conditions", "fail"}
REQUIRED_GATE_LABELS = (
    "Decision",
    "Evidence",
    "What this means",
    "Conditions or next step",
)
RESOURCE_LABELS = (
    "Sub-agent timing",
    "API-price comparison",
    "Human and other resources",
)
PLACEHOLDER_RE = re.compile(r"\b(?:TODO-FEASIBILITY|TBD|PLACEHOLDER)\b", re.IGNORECASE)
LABELED_FIELD_RE = re.compile(r"^\*\*([^*]+):\*\*\s*(.+)$", re.MULTILINE)
LEGACY_HEADINGS = {
    "task type",
    "variable verifiability",
    "either-way value",
    "data access",
    "base rate and power",
    "held-out validation",
    "time and cost",
    "inference",
    "human decision triggers",
}


def _decode_value(raw: str) -> str:
    value = raw.strip()
    if not value:
        return ""
    if value[0] == '"':
        try:
            decoded = json.loads(value)
        except json.JSONDecodeError as exc:
            raise ValueError(f"invalid quoted metadata value: {value}") from exc
        if not isinstance(decoded, str):
            raise ValueError(f"metadata values must be strings: {value}")
        return decoded.strip()
    if value[0] == "'":
        if len(value) < 2 or value[-1] != "'":
            raise ValueError(f"invalid quoted metadata value: {value}")
        return value[1:-1].replace("''", "'").strip()
    return value


def parse_source(text: str) -> tuple[dict[str, str], str]:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    if not lines or lines[0].strip() != "---":
        raise ValueError("source must begin with the feasibility-audit metadata block")
    try:
        end = next(index for index in range(1, len(lines)) if lines[index].strip() == "---")
    except StopIteration as exc:
        raise ValueError("metadata block has no closing --- line") from exc

    metadata: dict[str, str] = {}
    for line in lines[1:end]:
        if not line.strip() or line.lstrip().startswith("#"):
            continue
        key, separator, raw = line.partition(":")
        key = key.strip().lower()
        if not separator or not re.fullmatch(r"[a-z][a-z0-9_]*", key):
            raise ValueError(f"invalid metadata line: {line}")
        if key in metadata:
            raise ValueError(f"duplicate metadata key: {key}")
        metadata[key] = _decode_value(raw)

    missing = [key for key in REQUIRED_METADATA if not metadata.get(key)]
    if missing:
        raise ValueError("missing required metadata: " + ", ".join(missing))
    if PLACEHOLDER_RE.search(text):
        raise ValueError("source still contains an unresolved placeholder")
    recommendation = metadata["recommendation"].lower()
    if recommendation not in RECOMMENDATIONS:
        raise ValueError("recommendation must be go, go with modifications, or no-go")
    metadata["recommendation"] = recommendation
    try:
        date.fromisoformat(metadata["audit_date"])
    except ValueError as exc:
        raise ValueError("audit_date must use YYYY-MM-DD") from exc

    body = "\n".join(lines[end + 1 :]).strip()
    if not body:
        raise ValueError("source body is empty")
    return metadata, body


def _section_blocks(blocks, heading: str):
    start = next(
        index
        for index, block in enumerate(blocks)
        if block.kind == "heading" and block.level == 2 and block.text == heading
    )
    end = next(
        (
            index
            for index, block in enumerate(blocks[start + 1 :], start=start + 1)
            if block.kind == "heading" and block.level == 2
        ),
        len(blocks),
    )
    return blocks[start + 1 : end]


def _section_text(blocks, heading: str) -> str:
    return "\n".join(block.text for block in _section_blocks(blocks, heading) if block.text)


def _labeled_fields(section_text: str) -> dict[str, str]:
    fields: dict[str, str] = {}
    for match in LABELED_FIELD_RE.finditer(section_text):
        label = match.group(1).strip()
        fields.setdefault(label, match.group(2).strip())
    return fields


def validate_sections(blocks) -> None:
    if any(block.kind == "table" for block in blocks):
        raise ValueError("feasibility report must use prose sections, not a table")

    level_two = [block.text for block in blocks if block.kind == "heading" and block.level == 2]
    if level_two != list(REQUIRED_SECTIONS):
        missing = [section for section in REQUIRED_SECTIONS if section not in level_two]
        extra = [section for section in level_two if section not in REQUIRED_SECTIONS]
        detail = []
        if missing:
            detail.append("missing: " + "; ".join(missing))
        if extra:
            detail.append("unexpected: " + "; ".join(extra))
        if not detail:
            detail.append("required sections are misordered or duplicated")
        raise ValueError("invalid feasibility-report section headings (" + " | ".join(detail) + ")")

    legacy = [heading for heading in level_two if heading.strip().lower() in LEGACY_HEADINGS]
    if legacy:
        raise ValueError("legacy technical gate headings are not allowed: " + ", ".join(legacy))

    for section in REQUIRED_SECTIONS:
        if not _section_text(blocks, section).strip():
            raise ValueError(f"section contains no substantive content: {section}")

    for gate in GATES:
        fields = _labeled_fields(_section_text(blocks, gate.question))
        missing_labels = [label for label in REQUIRED_GATE_LABELS if not fields.get(label)]
        if missing_labels:
            raise ValueError(
                f"gate section '{gate.question}' is missing labeled content: "
                + ", ".join(missing_labels)
            )
        decision = fields["Decision"].splitlines()[0].strip().lower().rstrip(".")
        if decision not in GATE_DECISIONS:
            raise ValueError(
                f"gate section '{gate.question}' must decide pass, pass with conditions, or fail"
            )

    resource_fields = _labeled_fields(_section_text(blocks, GATES[5].question))
    missing_resources = [label for label in RESOURCE_LABELS if not resource_fields.get(label)]
    if missing_resources:
        raise ValueError(
            "time-and-resources gate is missing labeled content: " + ", ".join(missing_resources)
        )
    decision_fields = _labeled_fields(_section_text(blocks, GATES[7].question))
    if not decision_fields.get("Researcher decision needed"):
        raise ValueError("researcher-decision gate must state whether a researcher decision is needed")


def _add_metadata_line(doc: Document, label: str, value: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    label_run = paragraph.add_run(label + ": ")
    _format_run(label_run, "Calibri", 9.5, MUTED, bold=True)
    value_run = paragraph.add_run(value)
    _format_run(value_run, "Calibri", 9.5, DARK_BLUE)


def _add_rule(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), "CBD5E1")
    borders.append(bottom)
    paragraph._p.get_or_add_pPr().append(borders)


def build_document(metadata: dict[str, str], blocks) -> Document:
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = metadata["title"]
    doc.core_properties.subject = "ELARA Stage 03 feasibility audit"
    doc.core_properties.author = "ELARA"
    doc.core_properties.keywords = "feasibility audit, empirical legal research, research design"

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(4)
    kicker.paragraph_format.keep_with_next = True
    kicker_run = kicker.add_run("ELARA FEASIBILITY AUDIT")
    _format_run(kicker_run, "Calibri", 9, BLUE, bold=True)
    title = doc.add_paragraph(metadata["title"], style="ELARA Report Title")
    title.paragraph_format.keep_with_next = True
    subtitle = doc.add_paragraph(metadata["subtitle"], style="ELARA Report Subtitle")
    subtitle.paragraph_format.keep_with_next = True
    _add_metadata_line(doc, "Recommendation", metadata["recommendation"].title())
    _add_metadata_line(doc, "Audit date", metadata["audit_date"])
    _add_metadata_line(doc, "Report version", metadata["report_version"])
    _add_rule(doc)

    ordered_num: int | None = None
    bullet_num: int | None = None
    prior_kind = ""
    for block in blocks:
        if block.kind == "heading":
            style_level = min(max(block.level - 1, 1), 3)
            paragraph = doc.add_paragraph(style=f"Heading {style_level}")
            paragraph.paragraph_format.keep_with_next = True
            add_inline(paragraph, block.text)
        elif block.kind == "paragraph":
            paragraph = doc.add_paragraph()
            add_inline(paragraph, block.text)
        elif block.kind == "callout":
            paragraph = doc.add_paragraph(style="ELARA Callout")
            add_inline(paragraph, block.text)
            _set_paragraph_box(paragraph, fill=LIGHT_GRAY, border_color=BLUE)
        elif block.kind == "code":
            paragraph = doc.add_paragraph(style="ELARA Callout")
            run = paragraph.add_run(block.text)
            _format_run(run, "Consolas", 9, NAVY)
            _set_paragraph_box(paragraph, fill=LIGHT_GRAY, border_color="CBD5E1")
        elif block.kind in {"ordered", "bullet"}:
            if block.kind == "ordered" and prior_kind != "ordered":
                ordered_num = add_numbering_definition(doc, ordered=True)
            if block.kind == "bullet" and prior_kind != "bullet":
                bullet_num = add_numbering_definition(doc, ordered=False)
            paragraph = doc.add_paragraph()
            add_inline(paragraph, block.text)
            _apply_numbering(
                paragraph,
                ordered_num if block.kind == "ordered" else bullet_num,  # type: ignore[arg-type]
            )
        elif block.kind == "rule":
            _add_rule(doc)
        elif block.kind == "table":
            raise ValueError("feasibility report must use prose sections, not a table")
        prior_kind = block.kind
    return doc


def validate_docx_output(path: Path) -> None:
    document = Document(path)
    if document.tables:
        raise ValueError("generated feasibility report contains a table")
    headings = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.style is not None and paragraph.style.name == "Heading 1"
    ]
    if headings != list(REQUIRED_SECTIONS):
        raise ValueError("generated feasibility report has incorrect section headings")
    if any(PLACEHOLDER_RE.search(paragraph.text or "") for paragraph in document.paragraphs):
        raise ValueError("generated feasibility report contains an unresolved placeholder")
    if any(heading.lower() in LEGACY_HEADINGS for heading in headings):
        raise ValueError("generated feasibility report contains a legacy gate heading")


def build(source: Path, output: Path) -> None:
    if source.suffix.lower() != ".md":
        raise ValueError("source must be a .md file")
    output_format = output.suffix.lower()
    if output_format not in {".tex", ".docx"}:
        raise ValueError("output must be a .tex file, or .docx when Word was requested")
    if not source.is_file():
        raise ValueError(f"source does not exist: {source}")
    if output.exists():
        raise ValueError(f"refusing to overwrite existing artifact: {output}")
    metadata, body = parse_source(source.read_text(encoding="utf-8"))
    blocks = parse_blocks(body)
    validate_sections(blocks)
    output.parent.mkdir(parents=True, exist_ok=True)
    try:
        if output_format == ".tex":
            latex = render_latex_report(
                title=metadata["title"],
                subtitle=metadata["subtitle"],
                kicker="ELARA Feasibility Audit",
                metadata_rows=(
                    ("Recommendation", metadata["recommendation"].title()),
                    ("Audit date", metadata["audit_date"]),
                    ("Report version", metadata["report_version"]),
                ),
                blocks=blocks,
            )
            output.write_text(latex, encoding="utf-8", newline="\n")
            validate_latex_output(output, REQUIRED_SECTIONS, PLACEHOLDER_RE)
            if r"\begin{tabular" in latex:
                raise ValueError("generated feasibility report contains a table")
        else:
            document = build_document(metadata, blocks)
            document.save(output)
            validate_docx_output(output)
    except Exception:
        output.unlink(missing_ok=True)
        raise


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("source", type=Path, help="run-scoped Markdown build source")
    parser.add_argument(
        "output",
        type=Path,
        help="new versioned .tex artifact, or .docx after an explicit Word preference",
    )
    args = parser.parse_args()
    try:
        build(args.source.resolve(), args.output.resolve())
    except (OSError, ValueError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    print(f"Created {args.output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
