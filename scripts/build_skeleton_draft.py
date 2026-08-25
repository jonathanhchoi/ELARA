"""Validate and render an ELARA skeleton draft from one canonical Markdown source.

The source is an immutable, run-scoped planning artifact. The builder validates
its article structure, restrained prose, and provenance references, then creates
a new Word, LaTeX, or Markdown researcher-facing artifact and a run manifest.
Existing files are never overwritten.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import re
import sys
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.image.image import Image
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


CONTENT_WIDTH_DXA = 9360
NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "CBD5E1"
WHITE = "FFFFFF"

REQUIRED_METADATA = (
    "title",
    "subtitle",
    "output_format",
    "target_venue",
    "source_versions",
)
LEGACY_IGNORED_METADATA = {"target_length"}
SUPPORTED_FORMATS = {"docx", "tex", "md"}
WORD_TEMPLATE_REGISTRY = (
    Path(__file__).resolve().parents[1] / "workflow" / "templates" / "word" / "profiles.json"
)
REQUIRED_FIELDS = (
    "Section role",
    "Bare-bones content",
    "Source support",
    "Results presented",
    "Displays",
    "Author work",
    "Open questions",
)
LEGACY_IGNORED_FIELDS = {"Approximate length"}
TRACE_FIELDS = {"Bare-bones content", "Source support", "Results presented", "Displays"}
ALWAYS_TRACED_FIELDS = {"Source support", "Results presented", "Displays"}
ALLOWED_ROLES = {
    "abstract",
    "introduction",
    "background",
    "methods",
    "results",
    "robustness",
    "discussion",
    "limitations",
    "conclusion",
    "appendix",
    "other",
}
REQUIRED_TOP_LEVEL_ROLES = {
    "introduction",
    "methods",
    "results",
    "limitations",
    "conclusion",
}
EMPIRICAL_ROLES = {"methods", "results", "robustness"}
AUTHOR_TO_WRITE = "Author to write."
MAX_CONTENT_WORDS = 120
MAX_NONEMPIRICAL_CONTENT_WORDS = 35
PLACEHOLDER_RE = re.compile(
    r"\b(?:TODO-SKELETON|TODO|TBD|PLACEHOLDER|INSERT|XXX)\b|<[^>]+>",
    re.IGNORECASE,
)
SECTION_HEADING_RE = re.compile(r"^(#{2,6})\s+(.+?)\s*$")
FIELD_RE = re.compile(r"^\*\*([^*]+):\*\*\s*(.*?)\s*$")
PROJECT_REF_RE = re.compile(
    r"(?<![A-Za-z0-9_])(?P<path>project/[A-Za-z0-9_./-]+)#(?P<id>[A-Za-z0-9_.:-]+)"
)
BARE_PROJECT_REF_RE = re.compile(r"(?<![A-Za-z0-9_])project/[A-Za-z0-9_./-]+(?!#)")
NUMBER_RE = re.compile(r"(?<![A-Za-z])\b\d+(?:\.\d+)?%?\b")
DISPLAY_KINDS = {"table", "figure", "equation"}
DISPLAY_SUFFIXES = {
    "table": {".csv", ".tsv"},
    "figure": {".png", ".jpg", ".jpeg"},
    "equation": {".tex", ".txt"},
}
UNSAFE_TEX_RE = re.compile(
    r"\\(?:input|include|write|openout|read|usepackage|documentclass|begin\s*\{document\}|end\s*\{document\})\b",
    re.IGNORECASE,
)


@dataclass(frozen=True)
class Section:
    title: str
    level: int
    order: int
    fields: dict[str, str]

    @property
    def depth(self) -> int:
        return self.level - 2


@dataclass(frozen=True)
class Display:
    kind: str
    reference: str
    path: str
    artifact_id: str
    caption: str
    alt_text: str | None = None


def parse_displays(value: str) -> list[Display]:
    if value.strip().lower() == "none":
        return []
    displays: list[Display] = []
    for raw_spec in value.split(" || "):
        parts = [part.strip() for part in raw_spec.split("|", 3)]
        if len(parts) not in {3, 4} or not all(parts[:3]):
            raise ValueError(
                "Displays entries must use kind|project/path#artifact-id|caption or "
                "figure|project/path#artifact-id|caption|alt text, separated by ' || '"
            )
        kind, reference, caption = parts[:3]
        alt_text = parts[3] if len(parts) == 4 else None
        kind = kind.lower()
        if kind not in DISPLAY_KINDS:
            raise ValueError(f"unsupported display kind: {kind}")
        match = PROJECT_REF_RE.fullmatch(reference)
        if not match:
            raise ValueError(f"display reference must use project/path#artifact-id: {reference}")
        suffix = Path(match.group("path")).suffix.lower()
        if suffix not in DISPLAY_SUFFIXES[kind]:
            raise ValueError(f"{kind} display has unsupported file extension: {suffix or '(none)'}")
        if PLACEHOLDER_RE.search(caption):
            raise ValueError("display caption contains an unresolved placeholder")
        if alt_text and PLACEHOLDER_RE.search(alt_text):
            raise ValueError("figure alt text contains an unresolved placeholder")
        if alt_text and kind != "figure":
            raise ValueError("alt text is supported only for figure displays")
        displays.append(
            Display(kind, reference, match.group("path"), match.group("id"), caption, alt_text)
        )
    return displays


def _decode_value(raw: str) -> str:
    value = raw.strip()
    if not value:
        return ""
    if value.startswith('"'):
        try:
            parsed = json.loads(value)
        except json.JSONDecodeError as exc:
            raise ValueError(f"invalid quoted metadata value: {value}") from exc
        if not isinstance(parsed, str):
            raise ValueError("metadata values must be strings")
        return parsed.strip()
    if value.startswith("'"):
        if len(value) < 2 or not value.endswith("'"):
            raise ValueError(f"invalid quoted metadata value: {value}")
        return value[1:-1].replace("''", "'").strip()
    return value


def _source_version_paths(value: str) -> list[str]:
    paths = [item.strip() for item in value.split(";") if item.strip()]
    if not paths:
        raise ValueError("source_versions must name at least one verified project artifact")
    for item in paths:
        if not item.startswith("project/") or "#" in item or PLACEHOLDER_RE.search(item):
            raise ValueError(f"invalid source_versions entry: {item}")
    if len(paths) != len(set(paths)):
        raise ValueError("source_versions contains duplicate paths")
    return paths


def parse_source(text: str) -> tuple[dict[str, str], list[Section]]:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = normalized.split("\n")
    if not lines or lines[0].strip() != "---":
        raise ValueError("source must begin with a metadata block")
    try:
        metadata_end = next(i for i in range(1, len(lines)) if lines[i].strip() == "---")
    except StopIteration as exc:
        raise ValueError("metadata block has no closing --- line") from exc

    metadata: dict[str, str] = {}
    for line in lines[1:metadata_end]:
        if not line.strip() or line.lstrip().startswith("#"):
            continue
        key, separator, raw = line.partition(":")
        key = key.strip().lower()
        if not separator or not re.fullmatch(r"[a-z][a-z0-9_]*", key):
            raise ValueError(f"invalid metadata line: {line}")
        if key in metadata:
            raise ValueError(f"duplicate metadata key: {key}")
        metadata[key] = _decode_value(raw)

    missing = [key for key in REQUIRED_METADATA if not metadata.get(key)]
    if missing:
        raise ValueError("missing required metadata: " + ", ".join(missing))
    output_format = metadata["output_format"].lower().lstrip(".")
    if output_format not in SUPPORTED_FORMATS:
        raise ValueError("output_format must be docx, tex, or md")
    metadata["output_format"] = output_format
    _source_version_paths(metadata["source_versions"])
    if PLACEHOLDER_RE.search(normalized):
        raise ValueError("source contains an unresolved placeholder")
    for key in LEGACY_IGNORED_METADATA:
        metadata.pop(key, None)

    sections: list[Section] = []
    current_title = ""
    current_level = 0
    current_fields: dict[str, str] = {}

    def finish_section() -> None:
        nonlocal current_title, current_level, current_fields
        if not current_title:
            return
        missing_fields = [field for field in REQUIRED_FIELDS if not current_fields.get(field)]
        if missing_fields:
            raise ValueError(
                f"section {current_title} is missing required fields: {', '.join(missing_fields)}"
            )
        unknown_fields = sorted(
            set(current_fields) - set(REQUIRED_FIELDS) - LEGACY_IGNORED_FIELDS
        )
        if unknown_fields:
            raise ValueError(
                f"section {current_title} contains unknown fields: {', '.join(unknown_fields)}"
            )
        sections.append(
            Section(
                current_title,
                current_level,
                len(sections) + 1,
                {key: current_fields[key] for key in REQUIRED_FIELDS},
            )
        )
        current_title = ""
        current_level = 0
        current_fields = {}

    for line_number, line in enumerate(lines[metadata_end + 1 :], metadata_end + 2):
        if not line.strip():
            continue
        heading = SECTION_HEADING_RE.fullmatch(line.strip())
        if heading:
            finish_section()
            current_level = len(heading.group(1))
            current_title = heading.group(2).strip()
            continue
        if not current_title:
            raise ValueError(
                f"line {line_number} is outside a section; only structured section records are allowed"
            )
        field = FIELD_RE.fullmatch(line.strip())
        if not field:
            raise ValueError(
                f"line {line_number} is not a structured field; use the Bare-bones content field"
            )
        name, value = field.group(1).strip(), field.group(2).strip()
        if name in current_fields:
            raise ValueError(f"section {current_title} repeats field {name}")
        current_fields[name] = value
    finish_section()

    validate_sections(sections)
    return metadata, sections


def validate_sections(sections: list[Section]) -> None:
    top_level = [section for section in sections if section.depth == 0]
    if len(top_level) < 5:
        raise ValueError("skeleton draft must contain at least five top-level sections")
    titles = [section.title.casefold() for section in sections]
    if len(titles) != len(set(titles)):
        raise ValueError("section headings must be unique")

    previous_level = 1
    empirical_content_roles: set[str] = set()
    for section in sections:
        if section.level == 2:
            previous_level = 2
        elif section.level > previous_level + 1:
            raise ValueError(
                f"section {section.title} skips a Markdown heading level"
            )
        previous_level = section.level

        role = section.fields["Section role"].strip().lower()
        if role not in ALLOWED_ROLES:
            raise ValueError(
                f"section {section.title} has unsupported Section role {section.fields['Section role']}"
            )

        content = section.fields["Bare-bones content"].strip()
        word_count = len(re.findall(r"\b[\w'-]+\b", content))
        if word_count > MAX_CONTENT_WORDS:
            raise ValueError(
                f"section {section.title} Bare-bones content exceeds {MAX_CONTENT_WORDS} words"
            )
        if role not in EMPIRICAL_ROLES and content.casefold() != AUTHOR_TO_WRITE.casefold():
            if word_count > MAX_NONEMPIRICAL_CONTENT_WORDS:
                raise ValueError(
                    f"section {section.title} leaves too much non-empirical prose for the agent"
                )
            if len(re.findall(r"[.!?](?:\s|$)", content)) > 1:
                raise ValueError(
                    f"section {section.title} non-empirical content must be one sentence or less"
                )
        if role in EMPIRICAL_ROLES and content.casefold() != AUTHOR_TO_WRITE.casefold():
            empirical_content_roles.add(role)

        displays = parse_displays(section.fields["Displays"])
        results_value = section.fields["Results presented"].strip()
        if role in {"results", "robustness"}:
            if results_value.lower() == "none":
                raise ValueError(
                    f"section {section.title} must identify the results it presents"
                )
            if not displays:
                raise ValueError(
                    f"section {section.title} must present results through at least one table, figure, or equation"
                )

        section_refs = [
            match
            for field_name in TRACE_FIELDS
            for match in PROJECT_REF_RE.finditer(section.fields[field_name])
        ]
        for field_name in TRACE_FIELDS:
            value = section.fields[field_name]
            refs = list(PROJECT_REF_RE.finditer(value))
            value_without_refs = PROJECT_REF_RE.sub("", value)
            if BARE_PROJECT_REF_RE.search(value_without_refs):
                raise ValueError(
                    f"section {section.title} {field_name} artifact references require #ID"
                )
            if field_name in ALWAYS_TRACED_FIELDS and value.lower() != "none" and not refs:
                raise ValueError(
                    f"section {section.title} {field_name} must be 'none' or cite project/...#ID"
                )
            if NUMBER_RE.search(value) and not section_refs:
                raise ValueError(
                    f"section {section.title} {field_name} contains an untraced number"
                )

    top_roles = {section.fields["Section role"].strip().lower() for section in top_level}
    missing_roles = sorted(REQUIRED_TOP_LEVEL_ROLES - top_roles)
    if missing_roles:
        raise ValueError(
            "skeleton draft is not organizationally complete; missing top-level roles: "
            + ", ".join(missing_roles)
        )
    for required_role in ("methods", "results"):
        if required_role not in empirical_content_roles:
            raise ValueError(
                f"skeleton draft must include bare-bones {required_role} content"
            )


def validate_artifact_references(
    metadata: dict[str, str], sections: list[Section], project_root: Path
) -> None:
    repository_root = _repository_root(project_root)
    referenced = _source_version_paths(metadata["source_versions"])
    referenced.extend(
        match.group("path")
        for section in sections
        for value in section.fields.values()
        for match in PROJECT_REF_RE.finditer(value)
    )
    resolved = {
        item: _resolve_project_path(repository_root, item) for item in set(referenced)
    }
    missing = sorted(item for item, path in resolved.items() if not path.exists())
    if missing:
        raise ValueError("referenced project artifacts do not exist: " + ", ".join(missing))

    for section in sections:
        for display in parse_displays(section.fields["Displays"]):
            path = _resolve_project_path(repository_root, display.path)
            if display.kind == "table":
                _read_table(path)
            elif display.kind == "equation":
                _read_equation(path)


def _repository_root(project_root: Path) -> Path:
    return project_root.parent if project_root.name == "project" else project_root


def _resolve_project_path(repository_root: Path, value: str) -> Path:
    relative = Path(value)
    if relative.is_absolute() or not relative.parts or relative.parts[0] != "project":
        raise ValueError(f"artifact reference must remain under project/: {value}")
    project_directory = (repository_root / "project").resolve()
    resolved = (repository_root / relative).resolve()
    try:
        resolved.relative_to(project_directory)
    except ValueError as exc:
        raise ValueError(f"artifact reference escapes project/: {value}") from exc
    return resolved


def _read_table(path: Path) -> list[list[str]]:
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    with path.open(encoding="utf-8-sig", newline="") as handle:
        rows = [[cell.strip() for cell in row] for row in csv.reader(handle, delimiter=delimiter)]
    if not rows or not rows[0]:
        raise ValueError(f"display table is empty: {path}")
    width = len(rows[0])
    if any(len(row) != width for row in rows):
        raise ValueError(f"display table has inconsistent row widths: {path}")
    if len(rows) > 200 or width > 25:
        raise ValueError(f"display table is too large for a skeleton draft: {path}")
    return rows


def _read_equation(path: Path) -> str:
    value = path.read_text(encoding="utf-8").strip()
    if not value:
        raise ValueError(f"display equation is empty: {path}")
    if len(value) > 4000:
        raise ValueError(f"display equation is too long: {path}")
    if UNSAFE_TEX_RE.search(value):
        raise ValueError(f"display equation contains an unsafe LaTeX command: {path}")
    return value


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _set_style_font(style, size: float, color: str, *, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    for key in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{key}"), "Calibri")


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = False
    doc.settings.odd_and_even_pages_header_footer = False

    normal = doc.styles["Normal"]
    _set_style_font(normal, 10.5, "222222")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for name, size, before, after in (
        ("Title", 24, 0, 4),
        ("Subtitle", 11.5, 0, 14),
        ("Heading 1", 15, 14, 6),
        ("Heading 2", 12.5, 10, 4),
        ("Heading 3", 11.5, 8, 3),
        ("Heading 4", 10.5, 7, 3),
    ):
        style = doc.styles[name]
        _set_style_font(style, size, BLUE if name != "Title" else NAVY, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    table_width = OxmlElement("w:tblW")
    table_width.set(qn("w:w"), str(sum(widths)))
    table_width.set(qn("w:type"), "dxa")
    tbl_pr.insert(0, table_width)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), BORDER)
        borders.append(border)
    tbl_pr.append(borders)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_width = OxmlElement("w:tcW")
            tc_width.set(qn("w:w"), str(width))
            tc_width.set(qn("w:type"), "dxa")
            tc_pr.append(tc_width)
            margins = OxmlElement("w:tcMar")
            for side, value in (("top", 70), ("left", 100), ("bottom", 70), ("right", 100)):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
                margins.append(node)
            tc_pr.append(margins)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def _add_legacy_docx_displays(doc: Document, section: Section, repository_root: Path) -> None:
    for display in parse_displays(section.fields["Displays"]):
        path = _resolve_project_path(repository_root, display.path)
        label = display.kind.capitalize()
        caption = doc.add_paragraph()
        caption.paragraph_format.keep_with_next = True
        caption.add_run(f"{label}. ").bold = True
        caption.add_run(display.caption)
        source = doc.add_paragraph()
        source.paragraph_format.keep_with_next = True
        source_run = source.add_run(display.reference)
        source_run.italic = True
        source_run.font.size = Pt(8.5)
        source_run.font.color.rgb = RGBColor.from_string(MUTED)
        if display.kind == "table":
            rows = _read_table(path)
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            if rows:
                _repeat_header(table.rows[0])
            widths = [CONTENT_WIDTH_DXA // len(rows[0])] * len(rows[0])
            _set_table_geometry(table, widths)
            for row_index, (word_row, values) in enumerate(zip(table.rows, rows)):
                for cell, value in zip(word_row.cells, values):
                    cell.text = value
                    _set_cell_shading(cell, LIGHT_BLUE if row_index == 0 else WHITE)
                    if row_index == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
        elif display.kind == "figure":
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(path), width=Inches(6.2))
        else:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(_read_equation(path))
            run.font.name = "Cambria Math"
            run.font.size = Pt(11)


def build_legacy_docx(
    metadata: dict[str, str], sections: list[Section], output: Path, repository_root: Path
) -> None:
    doc = Document()
    _configure_document(doc)
    doc.core_properties.title = metadata["title"]
    doc.core_properties.subject = "ELARA Stage 17 article skeleton"
    doc.core_properties.author = "ELARA"
    fixed_time = datetime(2000, 1, 1, tzinfo=timezone.utc)
    doc.core_properties.created = fixed_time
    doc.core_properties.modified = fixed_time

    header = doc.sections[0].header.paragraphs[0]
    header.text = "ELARA  |  SKELETON DRAFT"
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MUTED)
    footer = doc.sections[0].footer.paragraphs[0]
    _add_page_number(footer)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    run = kicker.add_run("ELARA SKELETON DRAFT")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    doc.add_paragraph(metadata["title"], style="Title")
    doc.add_paragraph(metadata["subtitle"], style="Subtitle")

    callout = doc.add_table(rows=1, cols=1)
    _set_table_geometry(callout, [CONTENT_WIDTH_DXA])
    _set_cell_shading(callout.cell(0, 0), LIGHT_GRAY)
    paragraph = callout.cell(0, 0).paragraphs[0]
    paragraph.add_run("Authoring boundary. ").bold = True
    paragraph.add_run(
        "This draft supplies the complete article structure and only minimal methods and "
        "results language. The researcher writes the substantive prose."
    )

    for label, value in (
        ("Target venue", metadata["target_venue"]),
        ("Verified source set", metadata["source_versions"]),
    ):
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)

    doc.add_paragraph("Article structure", style="Heading 1")
    overview = doc.add_table(rows=1 + len(sections), cols=2)
    _repeat_header(overview.rows[0])
    for cell, value in zip(overview.rows[0].cells, ("Section", "Role")):
        _set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.add_run(value).bold = True
    for row, section in zip(overview.rows[1:], sections):
        cells = row.cells
        values = (
            f"{'  ' * section.depth}{section.title}",
            section.fields["Section role"],
        )
        for cell, value in zip(cells, values):
            cell.text = value
            _set_cell_shading(cell, WHITE)
    _set_table_geometry(overview, [7000, 2360])

    doc.add_page_break()
    doc.add_paragraph("Skeleton draft", style="Heading 1")
    for section in sections:
        heading_level = min(section.depth + 1, 4)
        doc.add_paragraph(section.title, style=f"Heading {heading_level}")
        content = doc.add_paragraph()
        content.paragraph_format.space_after = Pt(6)
        content_run = content.add_run(section.fields["Bare-bones content"])
        if section.fields["Bare-bones content"].casefold() == AUTHOR_TO_WRITE.casefold():
            content_run.italic = True
            content_run.font.color.rgb = RGBColor.from_string(MUTED)
        for field_name in (
            "Source support",
            "Results presented",
            "Author work",
            "Open questions",
        ):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.15)
            paragraph.paragraph_format.first_line_indent = Inches(-0.15)
            paragraph.add_run(f"{field_name}: ").bold = True
            paragraph.add_run(section.fields[field_name])
        _add_legacy_docx_displays(doc, section, repository_root)
    doc.save(output)


def _load_word_profiles() -> dict[str, dict[str, object]]:
    try:
        payload = json.loads(WORD_TEMPLATE_REGISTRY.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as exc:
        raise ValueError(f"cannot load Word template registry: {WORD_TEMPLATE_REGISTRY}") from exc
    profiles = payload.get("profiles")
    if not isinstance(profiles, dict) or not profiles:
        raise ValueError("Word template registry has no profiles")
    return profiles


def _is_yes(value: str) -> bool:
    return value.strip().casefold() in {"yes", "true", "1"}


def _is_jla_target(value: str) -> bool:
    normalized = re.sub(r"[^a-z]", "", value.casefold())
    return normalized in {"jla", "journaloflegalanalysis"}


def _word_template_info(
    metadata: dict[str, str], sections: list[Section]
) -> dict[str, object] | None:
    template_id = metadata.get("word_template", "").strip()
    if not template_id:
        return None
    if metadata["output_format"] != "docx":
        raise ValueError("word_template may be used only when output_format is docx")
    profiles = _load_word_profiles()
    if template_id not in profiles:
        raise ValueError(f"unsupported word_template: {template_id}")
    profile = dict(profiles[template_id])
    profile["template_id"] = template_id
    template_path = Path(__file__).resolve().parents[1] / str(profile["template_path"])
    if not template_path.is_file():
        raise ValueError(f"Word template does not exist: {profile['template_path']}")
    actual_hash = _sha256(template_path)
    expected_hash = str(profile.get("sha256", "")).casefold()
    if actual_hash != expected_hash:
        raise ValueError(
            f"Word template hash mismatch for {template_id}: expected {expected_hash}, got {actual_hash}"
        )
    profile["resolved_path"] = template_path

    fallback_approved = _is_yes(metadata.get("word_template_fallback_approved", ""))
    peer_reviewed = _is_yes(metadata.get("peer_reviewed", ""))
    is_jla_profile = template_id == "journal_of_legal_analysis_v1"
    custom_peer_venue = peer_reviewed and not _is_jla_target(metadata["target_venue"])
    if (is_jla_profile and not _is_jla_target(metadata["target_venue"])) or (
        custom_peer_venue and template_id == "law_review_v1"
    ):
        if not fallback_approved:
            raise ValueError(
                "a peer-reviewed outlet other than JLA requires an expressly approved Word "
                "template fallback; Stage 17 must first check the outlet's current official requirements"
            )
        if not metadata.get("venue_requirements_url") or not metadata.get(
            "venue_requirements_checked"
        ):
            raise ValueError(
                "an approved custom-outlet fallback requires venue_requirements_url and "
                "venue_requirements_checked"
            )

    if profile.get("figure_alt_text_required"):
        abstracts = [
            section.fields["Bare-bones content"]
            for section in sections
            if section.fields["Section role"].strip().casefold() == "abstract"
            and section.fields["Bare-bones content"].casefold() != AUTHOR_TO_WRITE.casefold()
        ]
        if abstracts and len(re.findall(r"\b[\w'-]+\b", abstracts[0])) > 100:
            raise ValueError("JLA abstract must not exceed 100 words")
        missing = [
            f"{section.title}: {display.caption}"
            for section in sections
            for display in parse_displays(section.fields["Displays"])
            if display.kind == "figure" and not display.alt_text
        ]
        if missing:
            raise ValueError("JLA figures require alt text: " + "; ".join(missing))
    return profile


def _load_template_without_prototype_comments(template_path: Path) -> Document:
    """Load a template and remove its demonstration comments without changing its package graph."""

    doc = Document(template_path)
    for comment in list(doc.part._comments_part.element):
        doc.part._comments_part.element.remove(comment)
    for expression in (
        ".//w:commentRangeStart",
        ".//w:commentRangeEnd",
        ".//w:commentReference",
    ):
        for node in list(doc._element.xpath(expression)):
            node.getparent().remove(node)
    return doc


def _replace_visible_text(paragraph, value: str) -> None:
    """Replace visible paragraph text while retaining fields and footnote references."""

    text_nodes = list(paragraph._p.iter(qn("w:t")))
    if text_nodes:
        text_nodes[0].text = value
        for node in text_nodes[1:]:
            node.text = ""
    else:
        paragraph.add_run(value)


def _clear_body_after(doc: Document, retained_elements: int) -> None:
    body = doc._body._element
    retained = 0
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        if retained < retained_elements:
            retained += 1
            continue
        body.remove(child)


def _request_field_updates(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def _replace_header_placeholder(doc: Document, running_title: str) -> None:
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            for paragraph in header.paragraphs:
                for node in paragraph._p.iter(qn("w:t")):
                    if node.text and "[RUNNING SHORT TITLE]" in node.text:
                        node.text = node.text.replace("[RUNNING SHORT TITLE]", running_title)


def _add_toc_field(doc: Document) -> None:
    heading = doc.add_paragraph("CONTENTS")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.paragraph_format.space_after = Pt(12)
    if heading.runs:
        heading.runs[0].bold = False
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def _visible_text(value: str) -> str:
    cleaned = PROJECT_REF_RE.sub("the verified project materials", value)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def _author_placeholder(value: str) -> str:
    text = _visible_text(value).strip()
    text = re.sub(r"\s+in the author(?:'s|’s) own prose\.?$", ".", text, flags=re.I)
    if len(text) > 180:
        text = text[:177].rstrip() + "..."
    return f"[Author: {text}]"


def _heading_labels(sections: list[Section], profile_id: str) -> dict[int, str]:
    labels: dict[int, str] = {}
    top_number = 0
    appendix_number = 0
    counters: list[int] = []
    in_appendix = False
    appendix_letter = ""
    for section in sections:
        depth = section.depth
        if depth == 0:
            counters = [0, 0, 0, 0]
            if section.fields["Section role"].strip().casefold() == "appendix":
                appendix_number += 1
                appendix_letter = chr(64 + appendix_number)
                in_appendix = True
                prefix = "APPENDIX" if profile_id == "law_review_v1" else "Appendix"
                labels[section.order] = f"{prefix} {appendix_letter}. {section.title}"
            else:
                top_number += 1
                in_appendix = False
                if profile_id == "law_review_v1":
                    roman = _roman(top_number)
                    labels[section.order] = f"{roman}. {section.title.upper()}"
                else:
                    labels[section.order] = f"{top_number}. {section.title}"
            continue
        while len(counters) <= depth:
            counters.append(0)
        counters[depth] += 1
        for index in range(depth + 1, len(counters)):
            counters[index] = 0
        if in_appendix:
            trail = ".".join(str(counters[index]) for index in range(1, depth + 1))
            punct = "." if profile_id == "law_review_v1" else ""
            labels[section.order] = f"{appendix_letter}.{trail}{punct} {section.title}"
        elif profile_id == "law_review_v1":
            if depth == 1:
                labels[section.order] = f"{chr(64 + counters[1])}. {section.title}"
            elif depth == 2:
                labels[section.order] = f"{counters[2]}. {section.title}"
            else:
                labels[section.order] = f"{counters[depth]}. {section.title}"
        else:
            number = [str(top_number)] + [str(counters[index]) for index in range(1, depth + 1)]
            labels[section.order] = f"{'.'.join(number)} {section.title}"
    return labels


def _roman(value: int) -> str:
    numerals = ((1000, "M"), (900, "CM"), (500, "D"), (400, "CD"), (100, "C"),
                (90, "XC"), (50, "L"), (40, "XL"), (10, "X"), (9, "IX"),
                (5, "V"), (4, "IV"), (1, "I"))
    result: list[str] = []
    for number, glyph in numerals:
        while value >= number:
            result.append(glyph)
            value -= number
    return "".join(result)


def _add_comment(doc: Document, paragraph, text: str) -> str:
    if not paragraph.runs:
        paragraph.add_run(" ")
    comment = doc.add_comment(
        runs=paragraph.runs,
        text=text,
        author="Research notes",
        initials="RN",
    )
    return str(comment.comment_id)


def _section_comment(section: Section) -> str:
    return "\n".join(
        (
            f"Source support: {section.fields['Source support']}",
            f"Results presented: {section.fields['Results presented']}",
            f"Open questions: {section.fields['Open questions']}",
        )
    )


def _content_width_dxa(doc: Document) -> int:
    section = doc.sections[0]
    width_inches = (
        section.page_width.inches - section.left_margin.inches - section.right_margin.inches
    )
    return max(1440, int(width_inches * 1440))


def _content_sensitive_widths(rows: list[list[str]], total_width: int) -> list[int]:
    weights = [
        max(6, min(48, max(len(row[index]) for row in rows) + 2))
        for index in range(len(rows[0]))
    ]
    minimum = min(900, total_width // len(weights))
    remaining = total_width - minimum * len(weights)
    weight_sum = sum(weights)
    widths = [minimum + int(remaining * weight / weight_sum) for weight in weights]
    widths[-1] += total_width - sum(widths)
    return widths


def _set_article_table_geometry(table, widths: list[int]) -> None:
    _set_table_geometry(table, widths)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        for border in borders:
            border.set(qn("w:color"), "000000")


def _embed_alt_text(shape, alt_text: str) -> None:
    shape._inline.docPr.set("descr", alt_text)
    for node in shape._inline.xpath(".//pic:cNvPr"):
        node.set("descr", alt_text)


def _add_venue_displays(
    doc: Document,
    section: Section,
    repository_root: Path,
    profile_id: str,
    counters: dict[str, int],
    comment_map: list[dict[str, object]],
) -> None:
    total_width = _content_width_dxa(doc)
    for display in parse_displays(section.fields["Displays"]):
        path = _resolve_project_path(repository_root, display.path)
        counters[display.kind] += 1
        number = counters[display.kind]
        label = display.kind.capitalize()
        if display.kind == "table":
            caption = doc.add_paragraph(f"Table {number}. {display.caption}", style="Caption")
            rows = _read_table(path)
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            _repeat_header(table.rows[0])
            _set_article_table_geometry(table, _content_sensitive_widths(rows, total_width))
            for row_index, (word_row, values) in enumerate(zip(table.rows, rows)):
                for cell, value in zip(word_row.cells, values):
                    cell.text = value
                    _set_cell_shading(cell, "D9D9D9" if row_index == 0 else WHITE)
                    if row_index == 0:
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
        elif display.kind == "figure":
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            natural = Image.from_file(str(path)).width.inches
            maximum = total_width / 1440
            shape = paragraph.add_run().add_picture(str(path), width=Inches(min(natural, maximum)))
            alt_text = display.alt_text or display.caption
            _embed_alt_text(shape, alt_text)
            caption = doc.add_paragraph(f"Figure {number}. {display.caption}", style="Caption")
            if profile_id == "journal_of_legal_analysis_v1":
                doc.add_paragraph(f"Alt text: {alt_text}")
        else:
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.add_run(_read_equation(path)).font.name = "Cambria Math"
            caption.add_run(f"\t({number})")
        comment_id = _add_comment(doc, caption, f"Display provenance: {display.reference}")
        comment_map.append(
            {
                "section": section.title,
                "anchor": f"{label} {number}",
                "comment_id": comment_id,
                "reference": display.reference,
            }
        )


def build_venue_docx(
    metadata: dict[str, str],
    sections: list[Section],
    output: Path,
    repository_root: Path,
    profile: dict[str, object],
) -> dict[str, object]:
    profile_id = str(profile["template_id"])
    doc = _load_template_without_prototype_comments(Path(profile["resolved_path"]))
    authors = metadata.get("authors", "").strip() or "[Author Name]"
    running_title = metadata.get("running_title", "").strip() or "[Running Short Title]"
    corresponding = (
        metadata.get("corresponding_author", "").strip()
        or "[Corresponding Author Contact Information]"
    )
    abstract_sections = [
        section
        for section in sections
        if section.fields["Section role"].strip().casefold() == "abstract"
    ]
    abstract_text = "[Abstract Contents]"
    if abstract_sections:
        candidate = abstract_sections[0].fields["Bare-bones content"]
        if candidate.casefold() != AUTHOR_TO_WRITE.casefold():
            abstract_text = _visible_text(candidate)
    body_sections = [section for section in sections if section not in abstract_sections]
    reference_sections = [
        section for section in body_sections if section.title.strip().casefold() == "references"
    ]
    body_sections = [section for section in body_sections if section not in reference_sections]
    comment_map: list[dict[str, object]] = []

    if profile_id == "law_review_v1":
        _replace_visible_text(doc.paragraphs[0], metadata["title"].upper())
        _replace_visible_text(doc.paragraphs[1], authors)
        _replace_visible_text(doc.paragraphs[2], abstract_text)
        _clear_body_after(doc, 3)
        if abstract_sections:
            comment_id = _add_comment(doc, doc.paragraphs[2], _section_comment(abstract_sections[0]))
            comment_map.append(
                {
                    "section": abstract_sections[0].title,
                    "anchor": "abstract",
                    "comment_id": comment_id,
                    "planning_fields": ["Source support", "Results presented", "Open questions"],
                }
            )
        _replace_header_placeholder(doc, running_title)
        doc.add_page_break()
        _add_toc_field(doc)
        doc.add_page_break()
    else:
        _replace_visible_text(doc.paragraphs[0], metadata["title"])
        _replace_visible_text(doc.paragraphs[1], authors)
        _replace_visible_text(
            doc.paragraphs[2], metadata.get("author_affiliations", "").strip() or "[Institutional Affiliation]"
        )
        _replace_visible_text(doc.paragraphs[3], f"Corresponding author: {corresponding}")
        _replace_visible_text(
            doc.paragraphs[4], "[Author: Add affiliations, funding, and acknowledgments.]"
        )
        _clear_body_after(doc, 5)
        doc.add_page_break()
        abstract_heading = doc.add_paragraph("Abstract", style="Heading 1")
        doc.add_paragraph(abstract_text)
        doc.add_paragraph("Keywords: [Keywords]")
        if abstract_sections:
            comment_id = _add_comment(doc, abstract_heading, _section_comment(abstract_sections[0]))
            comment_map.append(
                {
                    "section": abstract_sections[0].title,
                    "anchor": "abstract heading",
                    "comment_id": comment_id,
                    "planning_fields": ["Source support", "Results presented", "Open questions"],
                }
            )
        doc.add_page_break()

    _request_field_updates(doc)
    doc.core_properties.title = metadata["title"]
    doc.core_properties.subject = f"Manuscript skeleton for {metadata['target_venue']}"
    doc.core_properties.author = authors if not authors.startswith("[") else ""
    fixed_time = datetime(2000, 1, 1, tzinfo=timezone.utc)
    doc.core_properties.created = fixed_time
    doc.core_properties.modified = fixed_time

    labels = _heading_labels(body_sections, profile_id)
    counters = {"table": 0, "figure": 0, "equation": 0}
    for section in body_sections:
        heading = doc.add_paragraph(
            labels[section.order], style=f"Heading {min(section.depth + 1, 4)}"
        )
        comment_id = _add_comment(doc, heading, _section_comment(section))
        comment_map.append(
            {
                "section": section.title,
                "anchor": "heading",
                "comment_id": comment_id,
                "planning_fields": ["Source support", "Results presented", "Open questions"],
            }
        )
        content = section.fields["Bare-bones content"].strip()
        if content.casefold() == AUTHOR_TO_WRITE.casefold():
            doc.add_paragraph(_author_placeholder(section.fields["Author work"]), style="Author Placeholder")
        else:
            doc.add_paragraph(_visible_text(content))
            doc.add_paragraph(_author_placeholder(section.fields["Author work"]), style="Author Placeholder")
        _add_venue_displays(doc, section, repository_root, profile_id, counters, comment_map)

    if profile_id == "journal_of_legal_analysis_v1" or reference_sections:
        references = doc.add_paragraph("References", style="Heading 1")
        if reference_sections:
            for section in reference_sections:
                content = section.fields["Bare-bones content"]
                doc.add_paragraph(
                    _author_placeholder(section.fields["Author work"])
                    if content.casefold() == AUTHOR_TO_WRITE.casefold()
                    else _visible_text(content)
                )
                comment_id = _add_comment(doc, references, _section_comment(section))
                comment_map.append(
                    {
                        "section": section.title,
                        "anchor": "references heading",
                        "comment_id": comment_id,
                        "planning_fields": [
                            "Source support",
                            "Results presented",
                            "Open questions",
                        ],
                    }
                )
        elif profile_id == "journal_of_legal_analysis_v1":
            doc.add_paragraph("[Author: Add author-date references.]", style="Author Placeholder")

    doc.save(output)
    return {
        "template_id": profile_id,
        "template_path": str(profile["template_path"]),
        "template_sha256": str(profile["sha256"]),
        "requirements_authority": metadata.get("venue_requirements_authority")
        or profile.get("requirements_authority"),
        "requirements_checked": metadata.get("venue_requirements_checked")
        or profile.get("requirements_checked"),
        "requirements_url": metadata.get("venue_requirements_url")
        or profile.get("requirements_url"),
        "word_template_fallback_approved": _is_yes(
            metadata.get("word_template_fallback_approved", "")
        ),
        "comments_to_sections": comment_map,
    }


def build_docx(
    metadata: dict[str, str], sections: list[Section], output: Path, repository_root: Path
) -> dict[str, object] | None:
    profile = _word_template_info(metadata, sections)
    if profile is None:
        build_legacy_docx(metadata, sections, output, repository_root)
        return None
    return build_venue_docx(metadata, sections, output, repository_root, profile)


def _latex_escape(value: str) -> str:
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(replacements.get(char, char) for char in value)


def _latex_text_with_refs(value: str) -> str:
    pieces: list[str] = []
    cursor = 0
    for match in PROJECT_REF_RE.finditer(value):
        pieces.append(_latex_escape(value[cursor : match.start()]))
        pieces.append(r"\path{" + match.group(0) + "}")
        cursor = match.end()
    pieces.append(_latex_escape(value[cursor:]))
    return "".join(pieces)


def _latex_displays(display_value: str, repository_root: Path, output: Path) -> list[str]:
    lines: list[str] = []
    for display in parse_displays(display_value):
        path = _resolve_project_path(repository_root, display.path)
        lines.append(
            rf"\textbf{{{display.kind.capitalize()}.}} {_latex_escape(display.caption)}"
        )
        lines.append(rf"\textit{{Source:}} \path{{{display.reference}}}")
        if display.kind == "table":
            rows = _read_table(path)
            columns = "l" * len(rows[0])
            lines.extend([r"\begin{center}", r"\resizebox{\linewidth}{!}{%", rf"\begin{{tabular}}{{{columns}}}", r"\hline"])
            for index, row in enumerate(rows):
                content = " & ".join(_latex_escape(cell) for cell in row) + r" \\"
                if index == 0:
                    content = r"\textbf{" + (r"} & \textbf{".join(_latex_escape(cell) for cell in row)) + r"} \\ \hline"
                lines.append(content)
            lines.extend([r"\hline", r"\end{tabular}%", r"}", r"\end{center}"])
        elif display.kind == "figure":
            relative = Path(os.path.relpath(path, output.parent)).as_posix()
            lines.extend(
                [
                    r"\begin{center}",
                    rf"\includegraphics[width=0.92\linewidth]{{\detokenize{{{relative}}}}}",
                    r"\end{center}",
                ]
            )
        else:
            lines.extend([r"\[", _read_equation(path), r"\]"])
    return lines


def build_tex(
    metadata: dict[str, str], sections: list[Section], output: Path, repository_root: Path
) -> None:
    heading_commands = ("section", "subsection", "subsubsection", "paragraph")
    lines = [
        r"\documentclass[11pt]{article}",
        r"\usepackage[letterpaper,margin=1in]{geometry}",
        r"\usepackage[T1]{fontenc}",
        r"\usepackage{lmodern}",
        r"\usepackage{xcolor}",
        r"\usepackage{tabularx}",
        r"\usepackage{graphicx}",
        r"\usepackage{fancyhdr}",
        r"\usepackage[hidelinks]{hyperref}",
        r"\definecolor{ELARABlue}{HTML}{2E74B5}",
        r"\pagestyle{fancy}",
        r"\fancyhf{}",
        r"\lhead{\small\color{gray} ELARA | SKELETON DRAFT}",
        r"\rfoot{\thepage}",
        r"\fancypagestyle{plain}{\fancyhf{}\lhead{\small\color{gray} ELARA | SKELETON DRAFT}\rfoot{\thepage}}",
        r"\setlength{\parindent}{0pt}",
        r"\setlength{\parskip}{5pt}",
        r"\setlength{\emergencystretch}{3em}",
        r"\raggedright",
        r"\sloppy",
        rf"\title{{{_latex_escape(metadata['title'])}}}",
        rf"\author{{{_latex_escape(metadata['subtitle'])}}}",
        r"\date{}",
        r"\begin{document}",
        r"\maketitle",
        r"\begin{center}\fcolorbox{ELARABlue}{gray!8}{\parbox{0.88\linewidth}{\textbf{Authoring boundary.} This draft supplies the complete article structure and only minimal methods and results language. The researcher writes the substantive prose.}}\end{center}",
        r"\begin{itemize}",
        rf"\item \textbf{{Target venue:}} {_latex_escape(metadata['target_venue'])}",
        r"\item \textbf{Verified source set:}",
        r"\begin{itemize}",
        *(
            rf"\item \path{{{path}}}"
            for path in _source_version_paths(metadata["source_versions"])
        ),
        r"\end{itemize}",
        r"\end{itemize}",
        r"\section*{Article structure}",
        r"\begin{tabularx}{\linewidth}{@{}X p{0.22\linewidth}@{}}",
        r"\textbf{Section} & \textbf{Role} \\\hline",
    ]
    for section in sections:
        lines.append(
            f"{_latex_escape('  ' * section.depth + section.title)} & "
            f"{_latex_escape(section.fields['Section role'])} \\\\"
        )
    lines.extend([r"\end{tabularx}", r"\clearpage", r"\section*{Skeleton draft}"])
    for section in sections:
        command = heading_commands[min(section.depth, len(heading_commands) - 1)]
        lines.append(rf"\{command}*{{{_latex_escape(section.title)}}}")
        content = _latex_text_with_refs(section.fields["Bare-bones content"])
        if section.fields["Bare-bones content"].casefold() == AUTHOR_TO_WRITE.casefold():
            lines.append(rf"\textcolor{{gray}}{{\emph{{{content}}}}}")
        else:
            lines.append(content)
        lines.append(r"\begin{description}")
        for field_name in (
            "Source support",
            "Results presented",
            "Author work",
            "Open questions",
        ):
            lines.append(
                rf"\item[{_latex_escape(field_name)}] {_latex_text_with_refs(section.fields[field_name])}"
            )
        lines.append(r"\end{description}")
        lines.extend(_latex_displays(section.fields["Displays"], repository_root, output))
    lines.extend([r"\end{document}", ""])
    output.write_text("\n".join(lines), encoding="utf-8", newline="\n")


def _markdown_displays(display_value: str, repository_root: Path, output: Path) -> list[str]:
    lines: list[str] = []
    for display in parse_displays(display_value):
        path = _resolve_project_path(repository_root, display.path)
        lines.extend(
            [
                f"**{display.kind.capitalize()}.** {display.caption}",
                "",
                f"*Source: `{display.reference}`*",
                "",
            ]
        )
        if display.kind == "table":
            rows = _read_table(path)
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("|" + "|".join("---" for _ in rows[0]) + "|")
            lines.extend("| " + " | ".join(row) + " |" for row in rows[1:])
        elif display.kind == "figure":
            relative = Path(os.path.relpath(path, output.parent)).as_posix()
            lines.append(f"![{display.caption}]({relative})")
        else:
            lines.extend(["$$", _read_equation(path), "$$"])
        lines.append("")
    return lines


def build_markdown(
    metadata: dict[str, str], sections: list[Section], output: Path, repository_root: Path
) -> None:
    lines = [
        f"# {metadata['title']}",
        "",
        f"*{metadata['subtitle']}*",
        "",
        "> **Authoring boundary.** This draft supplies the complete article structure and only minimal methods and results language. The researcher writes the substantive prose.",
        "",
        f"- **Target venue:** {metadata['target_venue']}",
        f"- **Verified source set:** {metadata['source_versions']}",
        "",
        "## Article structure",
        "",
        "| Section | Role |",
        "|---|---|",
    ]
    for section in sections:
        lines.append(
            f"| {'&nbsp;' * (section.depth * 4)}{section.title} | "
            f"{section.fields['Section role']} |"
        )
    lines.extend(["", "## Skeleton draft", ""])
    for section in sections:
        level = min(3 + section.depth, 6)
        lines.append(f"{'#' * level} {section.title}")
        lines.append("")
        content = section.fields["Bare-bones content"]
        if content.casefold() == AUTHOR_TO_WRITE.casefold():
            lines.append(f"*{content}*")
        else:
            lines.append(content)
        lines.append("")
        for field_name in (
            "Source support",
            "Results presented",
            "Author work",
            "Open questions",
        ):
            lines.append(f"**{field_name}:** {section.fields[field_name]}")
            lines.append("")
        lines.extend(_markdown_displays(section.fields["Displays"], repository_root, output))
    output.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8", newline="\n")


def _validate_output(
    path: Path,
    output_format: str,
    sections: list[Section],
    metadata: dict[str, str],
) -> None:
    if output_format == "docx":
        reopened = Document(path)
        headings = {
            paragraph.text.strip()
            for paragraph in reopened.paragraphs
            if paragraph.style is not None and paragraph.style.name.startswith("Heading")
        }
        if metadata.get("word_template"):
            missing = [
                section.title
                for section in sections
                if section.fields["Section role"].strip().casefold() != "abstract"
                if not any(section.title.casefold() in heading.casefold() for heading in headings)
            ]
        else:
            missing = [section.title for section in sections if section.title not in headings]
        if missing:
            raise ValueError("generated Word artifact is missing sections: " + ", ".join(missing))
        has_table_display = any(
            display.kind == "table"
            for section in sections
            for display in parse_displays(section.fields["Displays"])
        )
        if has_table_display and not reopened.tables:
            raise ValueError("generated Word artifact is missing a table display")
        if metadata.get("word_template"):
            visible = "\n".join(
                [paragraph.text for paragraph in reopened.paragraphs]
                + [
                    cell.text
                    for table in reopened.tables
                    for row in table.rows
                    for cell in row.cells
                ]
            )
            if "ELARA" in visible:
                raise ValueError("venue-aware Word artifact contains visible ELARA branding")
            if "project/" in visible:
                raise ValueError("venue-aware Word artifact exposes a visible project path")
            if len(reopened.comments) < len(sections):
                raise ValueError("venue-aware Word artifact is missing planning comments")
    else:
        text = path.read_text(encoding="utf-8")
        missing = [section.title for section in sections if section.title not in text]
        if missing:
            raise ValueError("generated artifact is missing sections: " + ", ".join(missing))
        if PLACEHOLDER_RE.search(text):
            raise ValueError("generated artifact contains an unresolved placeholder")


def build(
    source: Path,
    output: Path,
    manifest: Path,
    project_root: Path,
) -> None:
    if source.suffix.lower() != ".md":
        raise ValueError("source must be a .md file")
    if not source.is_file():
        raise ValueError(f"source does not exist: {source}")
    for path in (output, manifest):
        if path.exists():
            raise ValueError(f"refusing to overwrite existing artifact: {path}")
    if manifest.suffix.lower() != ".json":
        raise ValueError("manifest must be a .json file")

    metadata, sections = parse_source(source.read_text(encoding="utf-8"))
    output_format = metadata["output_format"]
    if output.suffix.lower() != f".{output_format}":
        raise ValueError(
            f"output extension {output.suffix or '(none)'} does not match output_format {output_format}"
        )
    validate_artifact_references(metadata, sections, project_root)
    repository_root = _repository_root(project_root)

    for path in (output, manifest):
        path.parent.mkdir(parents=True, exist_ok=True)
    created: list[Path] = []
    try:
        word_build: dict[str, object] | None = None
        if output_format == "docx":
            word_build = build_docx(metadata, sections, output, repository_root)
        elif output_format == "tex":
            build_tex(metadata, sections, output, repository_root)
        else:
            build_markdown(metadata, sections, output, repository_root)
        created.append(output)
        _validate_output(output, output_format, sections, metadata)

        manifest_payload = {
            "schema_version": "2.0",
            "stage_id": "17-skeleton-draft",
            "source": {"path": source.as_posix(), "sha256": _sha256(source)},
            "output": {
                "path": output.as_posix(),
                "format": output_format,
                "sha256": _sha256(output),
            },
            "sections": [
                {
                    "title": section.title,
                    "level": section.level,
                    "order": section.order,
                    "role": section.fields["Section role"],
                    "planning_fields": dict(section.fields),
                }
                for section in sections
            ],
            "displays": [
                {
                    "section": section.title,
                    "kind": display.kind,
                    "reference": display.reference,
                    "caption": display.caption,
                    "alt_text": display.alt_text,
                }
                for section in sections
                for display in parse_displays(section.fields["Displays"])
            ],
            "verified_source_versions": _source_version_paths(metadata["source_versions"]),
            "word_template": (
                {
                    key: word_build.get(key)
                    for key in (
                        "template_id",
                        "template_path",
                        "template_sha256",
                        "requirements_authority",
                        "requirements_checked",
                        "requirements_url",
                        "word_template_fallback_approved",
                    )
                }
                if word_build
                else None
            ),
            "comments_to_sections": (
                word_build.get("comments_to_sections", []) if word_build else []
            ),
        }
        manifest.write_text(
            json.dumps(manifest_payload, indent=2, sort_keys=True) + "\n",
            encoding="utf-8",
            newline="\n",
        )
        created.append(manifest)
    except Exception:
        for path in reversed(created):
            path.unlink(missing_ok=True)
        raise


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("source", type=Path, help="immutable run-scoped Markdown source")
    parser.add_argument("output", type=Path, help="new .docx, .tex, or .md artifact")
    parser.add_argument("--manifest", type=Path, required=True, help="new run manifest JSON")
    parser.add_argument(
        "--project-root",
        type=Path,
        default=Path.cwd(),
        help="repository root containing project/",
    )
    args = parser.parse_args()
    try:
        build(
            args.source.resolve(),
            args.output.resolve(),
            args.manifest.resolve(),
            args.project_root.resolve(),
        )
    except (OSError, ValueError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    print(f"Created {args.output} and {args.manifest}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
