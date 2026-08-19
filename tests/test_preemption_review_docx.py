from __future__ import annotations

import sys
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from kit_context import resolve_test_root


ROOT = resolve_test_root(Path(__file__).resolve().parents[1])
SCRIPTS = ROOT / "scripts"
sys.path.insert(0, str(SCRIPTS))

from build_preemption_review import (  # noqa: E402
    ANNOTATED_MAP_SECTION,
    EXECUTIVE_SUMMARY_MAX_WORDS,
    REQUIRED_SECTIONS,
    build,
    parse_blocks,
    parse_source,
    validate_sections,
)
from workflow_lib import load_stages  # noqa: E402


SAMPLE_SOURCE = """---
title: "Preemption Review: Judicial Treatment of Synthetic Evidence"
subtitle: "How do federal courts evaluate authentication objections to synthetic evidence?"
verdict: "partially preempted"
recommended_disposition: "Reposition around authentication reasoning and the post-2024 period"
scoop_risk: "moderate"
review_date: "2026-08-19"
recheck_date: "2026-11-19"
---

## Executive summary

> **Bottom line:** The project is partially preempted because the doctrinal
thesis is occupied, but the proposed empirical comparison remains open.

**Closest threats:** Rivera is the closest work because it addresses the same
authentication problem and materially affects the verdict.

**Intended contribution:** The project would provide the first cross-court
measurement of authentication reasoning in post-2024 synthetic-evidence cases.

### Closest match: Rivera, *Authenticating Synthetic Media* (2025)

**What the work says:** Rivera asks how existing authentication doctrine applies
to synthetic media and argues for a reliability-focused doctrinal approach.

**Relevant scope and basis:** The article uses doctrinal analysis of reported
federal cases through 2024; it does not construct a cross-court opinion corpus.

**Preemption of the intended contribution:** Rivera occupies the project's
general doctrinal account of authentication problems, but not its proposed
measurement of variation across courts in post-2024 cases.

**What remains:** A cross-court empirical comparison using the proposed opinion
corpus remains unoccupied, provided the project drops its broad doctrinal claim.

**Evidence:** Source SRC-001 and claim-evidence entries CE-001 through CE-004.

**Remaining contribution:** The project can still measure how authentication
reasoning varies across courts after synthetic evidence became practical.

**Recommended disposition:** Reposition around authentication reasoning and
the post-2024 period, unless a working paper using the same corpus appears.

**Scoop risk and access gaps:** Risk is moderate; monitor two active researchers,
complete the HeinOnline search packet, and recheck by 2026-11-19.

## Annotated map of closest work

### Rivera, *Authenticating Synthetic Media* (2025)

**Publication status and venue:** Published law-review article.

**Question and thesis:** The article asks how existing authentication doctrine
applies to synthetic media and argues for a reliability-focused approach.

**Data and method:** Doctrinal analysis of reported cases.

**Relationship to this project:** It overlaps with the doctrinal question but
does not estimate court-level treatment using the proposed opinion corpus. See
[archived source](https://example.org/rivera).

## Verdict and flip conditions

> **Verdict:** Partially preempted. The doctrinal thesis is occupied, but the
proposed empirical comparison remains open.

- A working paper using the same opinion corpus would flip the verdict.
- A broader corpus can distinguish the project from the closest work.

## Positioning and lineage

**Honest contribution sentence:** The project measures how authentication
reasoning varies across courts after synthetic evidence became practically
available.

## Scoop risk

The risk is moderate because two active researchers have adjacent projects.

## Search methods and saturation evidence

1. Ran sixteen distinct queries across four routes.
2. Searched the closest authors and their coauthors.
3. Logged three saturation queries that produced no unseen close work.

| Route | Queries | Close works |
| --- | --- | --- |
| OpenAlex | 6 | 2 |
| Repositories | 5 | 1 |
| Open web | 5 | 0 |

## Access limitations and manual search packet

HeinOnline remained inaccessible. Run `"synthetic evidence" AND authentication`
and return any result published since 2024.

## Review date and recheck

Reviewed on 2026-08-19. Recheck by 2026-11-19 or sooner if a named researcher
posts a new working paper.
"""


class PreemptionReviewBuilderTests(unittest.TestCase):
    def test_wrapped_callouts_and_lists_remain_single_blocks(self) -> None:
        blocks = parse_blocks(
            "> **Verdict:** Partially preempted.\n"
            "The narrower empirical comparison remains open.\n\n"
            "- A working paper using the same corpus would\n"
            "  flip the verdict.\n\n"
            "1. Ran sixteen distinct queries across\n"
            "   four routes."
        )
        callout = next(block for block in blocks if block.kind == "callout")
        bullets = [block.text for block in blocks if block.kind == "bullet"]
        ordered = [block.text for block in blocks if block.kind == "ordered"]
        self.assertIn("comparison remains open", callout.text)
        self.assertIn("would flip the verdict", bullets[0])
        self.assertIn("four routes", ordered[0])

    def test_builds_structured_word_report_without_overwriting(self) -> None:
        with tempfile.TemporaryDirectory() as raw:
            root = Path(raw)
            source = root / "preemption_review_source.md"
            output = root / "preemption_review_v001.docx"
            source.write_text(SAMPLE_SOURCE, encoding="utf-8")

            build(source, output)

            self.assertTrue(output.is_file())
            self.assertGreater(output.stat().st_size, 10_000)
            document = Document(output)
            self.assertEqual(
                document.core_properties.title,
                "Preemption Review: Judicial Treatment of Synthetic Evidence",
            )
            self.assertGreaterEqual(len(document.tables), 2)
            self.assertEqual(document.tables[0].cell(0, 0).text, "Verdict")
            heading_text = {
                paragraph.text
                for paragraph in document.paragraphs
                if paragraph.style and paragraph.style.name == "Heading 1"
            }
            self.assertTrue(set(REQUIRED_SECTIONS) <= heading_text)
            ordered_headings = [
                paragraph.text
                for paragraph in document.paragraphs
                if paragraph.style and paragraph.style.name == "Heading 1"
            ]
            self.assertEqual(ordered_headings[:2], ["Executive summary", ANNOTATED_MAP_SECTION])
            annotated_map = next(
                paragraph for paragraph in document.paragraphs if paragraph.text == ANNOTATED_MAP_SECTION
            )
            self.assertTrue(annotated_map.paragraph_format.page_break_before)

            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
                relationships = archive.read("word/_rels/document.xml.rels").decode("utf-8")
                header_parts = [name for name in archive.namelist() if name.startswith("word/header")]
                footer_parts = [name for name in archive.namelist() if name.startswith("word/footer")]
            self.assertIn("w:numPr", document_xml)
            self.assertIn('w:w="9360"', document_xml)
            self.assertIn("https://example.org/rivera", relationships)
            self.assertEqual(header_parts, [])
            self.assertEqual(footer_parts, [])

            with self.assertRaisesRegex(ValueError, "refusing to overwrite"):
                build(source, output)

    def test_unresolved_template_marker_is_rejected(self) -> None:
        template = (ROOT / "workflow" / "templates" / "preemption_review_template.md").read_text(
            encoding="utf-8"
        )
        with self.assertRaisesRegex(ValueError, "unresolved placeholder"):
            parse_source(template)

    def test_executive_summary_is_required_and_capped(self) -> None:
        _metadata, body = parse_source(SAMPLE_SOURCE)
        blocks = parse_blocks(body)
        validate_sections(blocks)

        missing = SAMPLE_SOURCE.replace("## Executive summary", "## Overview", 1)
        _metadata, body = parse_source(missing)
        with self.assertRaisesRegex(ValueError, "Executive summary"):
            validate_sections(parse_blocks(body))

        oversized = SAMPLE_SOURCE.replace(
            "Risk is moderate; monitor two active researchers,",
            " ".join(["risk"] * (EXECUTIVE_SUMMARY_MAX_WORDS + 1)),
            1,
        )
        _metadata, body = parse_source(oversized)
        with self.assertRaisesRegex(ValueError, "exceeds 1200 words"):
            validate_sections(parse_blocks(body))

    def test_executive_summary_requires_a_complete_per_match_comparison(self) -> None:
        no_match = SAMPLE_SOURCE.replace(
            "### Closest match: Rivera, *Authenticating Synthetic Media* (2025)",
            "### Adjacent work: Rivera, *Authenticating Synthetic Media* (2025)",
            1,
        )
        _metadata, body = parse_source(no_match)
        with self.assertRaisesRegex(ValueError, "at least one 'Closest match:"):
            validate_sections(parse_blocks(body))

        no_preemption_analysis = SAMPLE_SOURCE.replace(
            "**Preemption of the intended contribution:**",
            "**General relationship:**",
            1,
        )
        _metadata, body = parse_source(no_preemption_analysis)
        with self.assertRaisesRegex(ValueError, "Preemption of the intended contribution"):
            validate_sections(parse_blocks(body))

        no_evidence = SAMPLE_SOURCE.replace(
            "**Evidence:** Source SRC-001 and claim-evidence entries CE-001 through CE-004.",
            "Source SRC-001 and claim-evidence entries CE-001 through CE-004.",
            1,
        )
        _metadata, body = parse_source(no_evidence)
        with self.assertRaisesRegex(ValueError, "Evidence"):
            validate_sections(parse_blocks(body))

        incomplete_second_match = SAMPLE_SOURCE.replace(
            "**Remaining contribution:**",
            "### Closest match: Chen, *Synthetic Proof in Federal Courts* (2026)\n\n"
            "**What the work says:** Chen reports that courts apply inconsistent "
            "authentication standards.\n\n"
            "**Relevant scope and basis:** The study codes published federal opinions "
            "from 2020 through 2025.\n\n"
            "**Preemption of the intended contribution:** It occupies the proposed "
            "cross-court measurement and most of the intended period.\n\n"
            "**What remains:** Only a later-period extension remains.\n\n"
            "**Remaining contribution:**",
            1,
        )
        _metadata, body = parse_source(incomplete_second_match)
        with self.assertRaisesRegex(
            ValueError,
            "Chen, .+ is missing required labeled content: Evidence",
        ):
            validate_sections(parse_blocks(body))


class PreemptionReviewStageContractTests(unittest.TestCase):
    def test_word_report_is_the_active_stage_02_artifact(self) -> None:
        stages = {meta["stage_id"]: (meta, body) for _, meta, body in load_stages(ROOT)}
        stage_two, body = stages["02-preemption-review"]
        self.assertIn(
            "project/artifacts/preemption_review_vNNN.docx",
            stage_two["declared_outputs"],
        )
        self.assertNotIn(
            "project/artifacts/preemption_review_vNNN.md",
            stage_two["declared_outputs"],
        )
        self.assertIn("project/runs/<run_id>/preemption_review_source.md", stage_two["declared_outputs"])
        self.assertIn("project/runs/<run_id>/rendered_preemption_review/", stage_two["declared_outputs"])
        self.assertIn("scripts/build_preemption_review.py", body)
        self.assertIn("decision-focused executive summary", body)
        self.assertIn("what the work actually says", body)
        self.assertIn("precise component of the intended contribution it occupies", body)
        self.assertIn("claim-evidence IDs or pinpoint pages", body)
        self.assertIn("1,200-word limit", body)
        self.assertIn("inspect every page at 100 percent zoom", body)
        for stage_id in ("03-feasibility-audit", "04-methods-design"):
            self.assertIn(
                "project/artifacts/preemption_review_vNNN.docx",
                stages[stage_id][0]["required_inputs"],
            )


if __name__ == "__main__":
    unittest.main()
