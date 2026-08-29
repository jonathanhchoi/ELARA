from __future__ import annotations

import json
import struct
import sys
import tempfile
import unittest
import zipfile
import zlib
import hashlib
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from build_skeleton_draft import build, parse_displays, parse_source  # noqa: E402


def sample_figure_png(width: int = 640, height: int = 360) -> bytes:
    """Return a dependency-free PNG with axes, intervals, and point estimates."""

    pixels = bytearray()
    estimates = [(150, 120), (270, 185), (390, 240), (510, 155)]
    for y in range(height):
        pixels.append(0)
        for x in range(width):
            color = (255, 255, 255)
            if (70 <= x <= 570 and y == 300) or (x == 70 and 45 <= y <= 300):
                color = (74, 85, 104)
            for point_x, point_y in estimates:
                if abs(x - point_x) <= 45 and abs(y - point_y) <= 2:
                    color = (75, 117, 181)
                if (x - point_x) ** 2 + (y - point_y) ** 2 <= 36:
                    color = (31, 78, 121)
            pixels.extend(color)

    def chunk(kind: bytes, payload: bytes) -> bytes:
        return (
            struct.pack(">I", len(payload))
            + kind
            + payload
            + struct.pack(">I", zlib.crc32(kind + payload) & 0xFFFFFFFF)
        )

    header = struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0)
    return b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", header) + chunk(
        b"IDAT", zlib.compress(bytes(pixels), 9)
    ) + chunk(b"IEND", b"")


def sample_source(output_format: str = "docx") -> str:
    return f'''---
title: "Court access and claim outcomes"
subtitle: "Skeleton draft"
output_format: "{output_format}"
target_venue: "Journal of Law and Empirical Analysis"
source_versions: "project/artifacts/preemption_review_v001.md; project/artifacts/methods_plan_v001.md; project/artifacts/results_v001.csv; project/artifacts/results_figure_v001.png; project/artifacts/estimating_equation_v001.tex; project/artifacts/robustness_v001.csv; project/DEVIATIONS.md"
---

## Introduction
**Section role:** introduction
**Bare-bones content:** Author to write.
**Source support:** project/artifacts/preemption_review_v001.md#CONTRIBUTION-01
**Displays:** none
**Author work:** State the research question, contribution, and thesis in the author's own prose.
**Open questions:** Whether to lead with court access or claim outcomes.

## Background and contribution
**Section role:** background
**Bare-bones content:** Author to write.
**Source support:** project/artifacts/preemption_review_v001.md#LITERATURE-01
**Displays:** none
**Author work:** Develop the literature position and legal context in the author's own prose.
**Open questions:** Whether the doctrinal discussion belongs here or after the results.

## Data and methods
**Section role:** methods
**Bare-bones content:** The study estimates filing outcomes in the approved court sample using the preregistered specification in project/artifacts/methods_plan_v001.md#METHOD-01.
**Source support:** project/artifacts/methods_plan_v001.md#METHOD-01; project/artifacts/methods_plan_v001.md#ESTIMAND-01
**Displays:** equation|project/artifacts/estimating_equation_v001.tex#EQ-01|where Y_ic is the outcome for claim i in court c, T_ic is the treatment indicator, X_ic collects the covariates, delta_c is the court fixed effect, and epsilon_ic is the idiosyncratic error
**Author work:** Explain the design choices and identifying assumptions in the author's own prose.
**Open questions:** Whether to move implementation details to an appendix.

### Validation
**Section role:** methods
**Bare-bones content:** The held-out comparison met the approved validation threshold in project/artifacts/methods_plan_v001.md#VALIDATION-01.
**Source support:** project/artifacts/methods_plan_v001.md#VALIDATION-01
**Displays:** none
**Author work:** Explain the remaining measurement concern.
**Open questions:** Whether validation details belong in the main text.

## Results
**Section role:** results
**Bare-bones content:** The primary estimate is positive, the secondary estimate is null, and uncertainty is reported in project/artifacts/results_v001.csv#RESULTS-ALL.
**Source support:** project/artifacts/results_v001.csv#RESULTS-ALL
**Displays:** table|project/artifacts/results_v001.csv#TABLE-PRIMARY|Complete estimates for the primary, secondary, and subgroup analyses with standard errors and sample sizes || figure|project/artifacts/results_figure_v001.png#FIGURE-PRIMARY|Point estimates and 95 percent confidence intervals for every preregistered analysis
**Author work:** Interpret the complete findings and connect them to the argument in the author's own prose.
**Open questions:** Which result should receive the most attention in the introduction.

### Robustness and deviations
**Section role:** robustness
**Bare-bones content:** The complete specification checks and the fragile subgroup result appear in project/artifacts/robustness_v001.csv#ROBUSTNESS-ALL.
**Source support:** project/artifacts/robustness_v001.csv#ROBUSTNESS-ALL; project/DEVIATIONS.md#DEV-01
**Displays:** table|project/artifacts/robustness_v001.csv#TABLE-ROBUSTNESS|All robustness specifications, including null and fragile findings, with standard errors and sample sizes
**Author work:** Explain how the checks and deviation affect the conclusions.
**Open questions:** Whether the full table belongs in the main text or appendix.

## Limitations
**Section role:** limitations
**Bare-bones content:** Author to write.
**Source support:** project/DEVIATIONS.md#LIMITATION-01
**Displays:** none
**Author work:** Develop the limits on inference and generalization in the author's own prose.
**Open questions:** How much space to give measurement and coverage limits.

## Conclusion
**Section role:** conclusion
**Bare-bones content:** Author to write.
**Source support:** project/artifacts/preemption_review_v001.md#CONTRIBUTION-01
**Displays:** none
**Author work:** State the implications and conclusion in the author's own prose.
**Open questions:** none
'''


def legacy_sample_source(output_format: str = "docx") -> str:
    lines: list[str] = []
    for line in sample_source(output_format).splitlines():
        lines.append(line)
        if line.startswith("target_venue:"):
            lines.append('target_length: "10,000 words"')
        elif line.startswith("**Source support:**"):
            lines.append(
                "**Results presented:** project/artifacts/results_v001.csv#R-H01"
            )
        elif line.startswith("**Open questions:**"):
            lines.append("**Approximate length:** 500 words")
    return "\n".join(lines) + "\n"


def venue_source(profile: str, *, include_authors: bool = True) -> str:
    text = sample_source("docx")
    additions = [f'word_template: "{profile}"']
    if include_authors:
        additions.extend(
            [
                'authors: "Alice Example; Bob Example"',
                'running_title: "Court Access"',
                'corresponding_author: "Alice Example, alice@example.edu"',
            ]
        )
    text = text.replace("source_versions:", "\n".join(additions) + "\nsource_versions:", 1)
    if profile == "journal_of_legal_analysis_v1":
        text = text.replace(
            'target_venue: "Journal of Law and Empirical Analysis"',
            'target_venue: "Journal of Legal Analysis"',
        )
        text = text.replace(
            "Point estimates and 95 percent confidence intervals for every preregistered analysis",
            "Point estimates and 95 percent confidence intervals for every preregistered analysis|Dot-and-whisker plot of four estimates with horizontal 95 percent confidence intervals",
            1,
        )
    return text


def rendered_text(path: Path, output_format: str) -> str:
    if output_format != "docx":
        return path.read_text(encoding="utf-8")
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    parts.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return "\n".join(parts)


class SkeletonDraftBuilderTests(unittest.TestCase):
    def _project(self, root: Path) -> None:
        text_files = {
            "project/artifacts/preemption_review_v001.md": "verified contribution and literature\n",
            "project/artifacts/methods_plan_v001.md": "verified methods and validation\n",
            "project/artifacts/estimating_equation_v001.tex": r"Y_{ic} = \alpha + \beta T_{ic} + X_{ic}'\gamma + \delta_c + \varepsilon_{ic}",
            "project/DEVIATIONS.md": "verified deviations and limitations\n",
        }
        for relative, value in text_files.items():
            path = root / relative
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_text(value, encoding="utf-8")
        for relative, value in {
            "project/artifacts/results_v001.csv": "analysis,estimate,se,n\nprimary,0.12,0.03,400\nsecondary,0.01,0.04,400\nsubgroup,0.08,0.06,90\n",
            "project/artifacts/robustness_v001.csv": "specification,estimate,se,n\nbaseline,0.12,0.03,400\nalternate,0.11,0.04,400\nsubgroup,0.08,0.06,90\n",
        }.items():
            path = root / relative
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_text(value, encoding="utf-8")
        figure = root / "project/artifacts/results_figure_v001.png"
        figure.parent.mkdir(parents=True, exist_ok=True)
        figure.write_bytes(sample_figure_png())

    def _build(self, root: Path, output_format: str) -> tuple[Path, Path, Path]:
        self._project(root)
        source = root / f"source_{output_format}.md"
        output = root / f"skeleton_draft_v001.{output_format}"
        manifest = root / f"manifest_{output_format}.json"
        source.write_text(sample_source(output_format), encoding="utf-8")
        build(source, output, manifest, root)
        return source, output, manifest

    def test_builds_all_supported_formats_with_complete_displays(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            for output_format in ("docx", "tex", "md"):
                source, output, manifest = self._build(root / output_format, output_format)
                self.assertTrue(source.is_file())
                self.assertTrue(output.is_file())
                payload = json.loads(manifest.read_text(encoding="utf-8"))
                self.assertEqual(payload["schema_version"], "2.0")
                self.assertEqual(payload["output"]["format"], output_format)
                self.assertEqual(payload["sections"][0]["title"], "Introduction")
                self.assertEqual(payload["sections"][3]["level"], 3)
                self.assertEqual(len(payload["displays"]), 4)
                self.assertNotIn("crosswalk", payload)
                self.assertNotIn(
                    "Results presented", payload["sections"][0]["planning_fields"]
                )
                text = rendered_text(output, output_format)
                self.assertNotIn("Target length", text)
                self.assertNotIn("Approximate length", text)
                self.assertNotIn("Results presented", text)

    def _build_venue(self, root: Path, profile: str, *, include_authors: bool = True):
        self._project(root)
        source = root / "source.md"
        output = root / "skeleton.docx"
        manifest = root / "manifest.json"
        source.write_text(
            venue_source(profile, include_authors=include_authors), encoding="utf-8"
        )
        build(source, output, manifest, root)
        return output, json.loads(manifest.read_text(encoding="utf-8"))

    def test_bundled_templates_match_approved_hashes(self) -> None:
        expected = {
            "law_review_v1.docx": "727a3747ef582cb172159c4aa42da5c6cca2d2383ae980f8059fa645f0cad8e8",
            "journal_of_legal_analysis_v1.docx": "8a5385ccfcc70d6bdbff67421828342dcfab3c674633ea35938580c5220f91e3",
        }
        template_root = ROOT / "workflow/templates/word"
        for name, digest in expected.items():
            self.assertEqual(hashlib.sha256((template_root / name).read_bytes()).hexdigest(), digest)

        with zipfile.ZipFile(template_root / "journal_of_legal_analysis_v1.docx") as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")
            footnotes_xml = archive.read("word/footnotes.xml").decode("utf-8")
        self.assertIn("footnoteReference", document_xml)
        self.assertIn("Substantive Footnote Contents", footnotes_xml)

    def test_law_review_template_output_has_article_geometry_toc_comments_and_footnote(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output, payload = self._build_venue(Path(directory), "law_review_v1")
            document = Document(output)
            section = document.sections[0]
            self.assertAlmostEqual(section.left_margin.inches, 1.5, places=2)
            self.assertAlmostEqual(section.right_margin.inches, 1.5, places=2)
            self.assertAlmostEqual(section.top_margin.inches, 1.0, places=2)
            self.assertAlmostEqual(section.bottom_margin.inches, 1.0, places=2)
            self.assertEqual(document.styles["Normal"].font.name, "Century Schoolbook")
            self.assertEqual(document.paragraphs[0].text, "COURT ACCESS AND CLAIM OUTCOMES")
            self.assertIn("Alice Example; Bob Example", document.paragraphs[1].text)
            visible = rendered_text(output, "docx")
            self.assertNotIn("ELARA", visible)
            self.assertNotIn("project/", visible)
            self.assertNotIn("Article structure", visible)
            self.assertIn("[Author:", visible)
            headings = [
                paragraph.text
                for paragraph in document.paragraphs
                if paragraph.style and paragraph.style.name.startswith("Heading")
            ]
            self.assertIn("I. INTRODUCTION", headings)
            self.assertIn("A. Validation", headings)
            sections_with_questions = [
                section
                for section in payload["sections"]
                if section["planning_fields"]["Open questions"].strip().lower() != "none"
            ]
            self.assertEqual(len(document.comments), len(sections_with_questions))
            self.assertLess(len(document.comments), len(payload["sections"]))
            self.assertTrue(
                all(
                    entry["planning_fields"] == ["Open questions"]
                    for entry in payload["comments_to_sections"]
                )
            )
            self.assertNotIn(
                "Conclusion",
                {entry["section"] for entry in payload["comments_to_sections"]},
            )
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            equation_index = next(
                index for index, text in enumerate(paragraphs) if "Y_{ic}" in text
            )
            self.assertIn("court fixed effect", paragraphs[equation_index + 1])
            self.assertEqual(payload["word_template"]["template_id"], "law_review_v1")
            self.assertEqual(
                payload["word_template"]["template_sha256"],
                "727a3747ef582cb172159c4aa42da5c6cca2d2383ae980f8059fa645f0cad8e8",
            )
            self.assertEqual(len(payload["comments_to_sections"]), len(document.comments))
            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
                footnotes_xml = archive.read("word/footnotes.xml").decode("utf-8")
                styles_xml = archive.read("word/styles.xml").decode("utf-8")
                comments_xml = archive.read("word/comments.xml").decode("utf-8")
            self.assertIn("Open questions:", comments_xml)
            self.assertNotIn("Source support", comments_xml)
            self.assertNotIn("Results presented", comments_xml)
            self.assertNotIn("Display provenance", comments_xml)
            self.assertIn('<w:sz w:val="22"', styles_xml)
            self.assertIn('TOC \\o "1-3"', document_xml)
            self.assertIn("footnoteReference", document_xml)
            self.assertIn("Information; Acknowledgments", footnotes_xml)
            self.assertIn('w:tblHeader w:val="true"', document_xml)
            self.assertIn(
                'descr="Point estimates and 95 percent confidence intervals for every preregistered analysis"',
                document_xml,
            )

    def test_jla_template_output_has_required_layout_alt_text_and_references(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output, payload = self._build_venue(
                Path(directory), "journal_of_legal_analysis_v1"
            )
            document = Document(output)
            section = document.sections[0]
            for margin in (
                section.left_margin,
                section.right_margin,
                section.top_margin,
                section.bottom_margin,
            ):
                self.assertAlmostEqual(margin.inches, 1.25, places=2)
            self.assertEqual(document.styles["Normal"].font.name, "Times New Roman")
            self.assertAlmostEqual(document.styles["Normal"].font.size.pt, 12.0, places=1)
            self.assertAlmostEqual(document.styles["Normal"].paragraph_format.line_spacing, 2.0)
            visible = rendered_text(output, "docx")
            self.assertIn("Corresponding author: Alice Example, alice@example.edu", visible)
            self.assertIn("Alt text: Dot-and-whisker plot", visible)
            self.assertIn("References", visible)
            self.assertNotIn("ELARA", visible)
            self.assertNotIn("project/", visible)
            headings = [
                paragraph.text
                for paragraph in document.paragraphs
                if paragraph.style and paragraph.style.name.startswith("Heading")
            ]
            self.assertIn("1. Introduction", headings)
            self.assertIn("3.1 Validation", headings)
            self.assertEqual(
                payload["word_template"]["template_id"],
                "journal_of_legal_analysis_v1",
            )
            figure = next(display for display in payload["displays"] if display["kind"] == "figure")
            self.assertTrue(figure["alt_text"].startswith("Dot-and-whisker"))
            self.assertIn("Source support", payload["sections"][0]["planning_fields"])
            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
            self.assertIn("Dot-and-whisker plot", document_xml)
            self.assertNotIn(" TOC ", document_xml)

    def test_author_owned_defaults_are_explicit_placeholders(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            output, _ = self._build_venue(
                Path(directory), "law_review_v1", include_authors=False
            )
            document = Document(output)
            self.assertIn("[Author Name]", document.paragraphs[1].text)
            self.assertIn("[Abstract Contents]", rendered_text(output, "docx"))

    def test_rejects_invalid_template_missing_jla_alt_text_and_silent_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            self._project(root)
            source = root / "source.md"
            output = root / "out.docx"
            manifest = root / "manifest.json"
            source.write_text(
                venue_source("law_review_v1").replace("law_review_v1", "missing_v1", 1),
                encoding="utf-8",
            )
            with self.assertRaisesRegex(ValueError, "unsupported word_template"):
                build(source, output, manifest, root)
            source.write_text(
                venue_source("journal_of_legal_analysis_v1").replace(
                    "|Dot-and-whisker plot of four estimates with horizontal 95 percent confidence intervals",
                    "",
                    1,
                ),
                encoding="utf-8",
            )
            with self.assertRaisesRegex(ValueError, "require alt text"):
                build(source, output, manifest, root)
            custom = venue_source("journal_of_legal_analysis_v1").replace(
                'target_venue: "Journal of Legal Analysis"',
                'target_venue: "Custom Peer Review Journal"',
            )
            source.write_text(custom, encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "expressly approved Word template fallback"):
                build(source, output, manifest, root)
            source.write_text(
                custom.replace(
                    "word_template:",
                    'word_template_fallback_approved: "yes"\nword_template:',
                    1,
                ),
                encoding="utf-8",
            )
            with self.assertRaisesRegex(ValueError, "venue_requirements_url"):
                build(source, output, manifest, root)

    def test_legacy_fields_are_accepted_but_not_rendered(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            for output_format in ("docx", "tex", "md"):
                format_root = root / output_format
                self._project(format_root)
                source = format_root / "legacy_source.md"
                output = format_root / f"skeleton_draft_v001.{output_format}"
                manifest = format_root / "manifest.json"
                source.write_text(legacy_sample_source(output_format), encoding="utf-8")
                metadata, sections = parse_source(source.read_text(encoding="utf-8"))
                self.assertNotIn("target_length", metadata)
                self.assertTrue(
                    all("Approximate length" not in section.fields for section in sections)
                )
                self.assertTrue(
                    all("Results presented" not in section.fields for section in sections)
                )
                build(source, output, manifest, format_root)
                text = rendered_text(output, output_format)
                self.assertNotIn("Target length", text)
                self.assertNotIn("Approximate length", text)
                self.assertNotIn("Results presented", text)
                self.assertNotIn("10,000 words", text)
                self.assertNotIn("500 words", text)

    def test_word_uses_descriptive_headings_and_renders_all_display_types(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            _, output, _ = self._build(Path(directory), "docx")
            document = Document(output)
            headings = [
                paragraph.text
                for paragraph in document.paragraphs
                if paragraph.style and paragraph.style.name.startswith("Heading")
            ]
            self.assertIn("Introduction", headings)
            self.assertIn("Validation", headings)
            self.assertFalse(any("[S0" in heading for heading in headings))
            self.assertGreaterEqual(len(document.tables), 3)
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            equation_index = next(
                index for index, text in enumerate(paragraphs) if "Y_{ic}" in text
            )
            self.assertTrue(paragraphs[equation_index + 1].startswith("Equation. "))
            with zipfile.ZipFile(output) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
                footer_xml = "".join(
                    archive.read(name).decode("utf-8")
                    for name in archive.namelist()
                    if name.startswith("word/footer") and name.endswith(".xml")
                )
                media = [name for name in archive.namelist() if name.startswith("word/media/")]
            self.assertIn('w:type="fixed"', document_xml)
            self.assertIn(" PAGE ", footer_xml)
            self.assertTrue(media)

    def test_latex_and_markdown_include_tables_figures_equations_and_captions(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            _, tex_output, _ = self._build(root / "tex", "tex")
            tex = tex_output.read_text(encoding="utf-8")
            self.assertIn(r"\section*{Introduction}", tex)
            self.assertIn(r"\subsection*{Validation}", tex)
            self.assertIn(r"\includegraphics", tex)
            self.assertIn(r"Y_{ic}", tex)
            self.assertIn("Complete estimates for the primary", tex)
            self.assertLess(tex.index("Y_{ic}"), tex.index("court fixed effect"))
            _, md_output, _ = self._build(root / "md", "md")
            md = md_output.read_text(encoding="utf-8")
            self.assertIn("| analysis | estimate | se | n |", md)
            self.assertIn("![Point estimates", md)
            self.assertIn("$$", md)
            self.assertLess(md.index("$$"), md.index("**Equation.**"))

    def test_rejects_unsupported_format_and_extension_mismatch(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            self._project(root)
            source = root / "source.md"
            source.write_text(sample_source("pdf"), encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "output_format"):
                build(source, root / "out.pdf", root / "manifest.json", root)
            source.write_text(sample_source("tex"), encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "does not match"):
                build(source, root / "out.md", root / "manifest.json", root)

    def test_rejects_missing_field_placeholder_and_heading_jump(self) -> None:
        missing = sample_source().replace("**Open questions:** Whether to lead with court access or claim outcomes.\n", "", 1)
        with self.assertRaisesRegex(ValueError, "missing required fields"):
            parse_source(missing)
        with self.assertRaisesRegex(ValueError, "unresolved placeholder"):
            parse_source(sample_source().replace("Skeleton draft", "TODO-SKELETON", 1))
        with self.assertRaisesRegex(ValueError, "skips a Markdown heading level"):
            parse_source(sample_source().replace("### Validation", "#### Validation", 1))

    def test_requires_complete_structure_and_minimal_methods_and_results_prose(self) -> None:
        missing_conclusion = sample_source().split("\n## Conclusion\n", 1)[0] + "\n"
        with self.assertRaisesRegex(ValueError, "missing top-level roles: conclusion"):
            parse_source(missing_conclusion)
        no_results_content = sample_source().replace(
            "The primary estimate is positive, the secondary estimate is null, and uncertainty is reported in project/artifacts/results_v001.csv#RESULTS-ALL.",
            "Author to write.",
        )
        with self.assertRaisesRegex(ValueError, "bare-bones results content"):
            parse_source(no_results_content)
        excess = sample_source().replace(
            "**Bare-bones content:** Author to write.",
            "**Bare-bones content:** " + " ".join(["word"] * 36) + ".",
            1,
        )
        with self.assertRaisesRegex(ValueError, "too much non-empirical prose"):
            parse_source(excess)

    def test_requires_displayed_result_sections_and_traced_references(self) -> None:
        no_display = sample_source().replace(
            "**Displays:** table|project/artifacts/results_v001.csv#TABLE-PRIMARY|Complete estimates for the primary, secondary, and subgroup analyses with standard errors and sample sizes || figure|project/artifacts/results_figure_v001.png#FIGURE-PRIMARY|Point estimates and 95 percent confidence intervals for every preregistered analysis",
            "**Displays:** none",
        )
        with self.assertRaisesRegex(ValueError, "at least one table, figure, or equation"):
            parse_source(no_display)
        partly_bare = sample_source().replace(
            "**Source support:** project/artifacts/results_v001.csv#RESULTS-ALL",
            "**Source support:** project/artifacts/results_v001.csv#RESULTS-ALL; project/DEVIATIONS.md",
        )
        with self.assertRaisesRegex(ValueError, "artifact references require #ID"):
            parse_source(partly_bare)

    def test_rejects_bad_display_specs_and_missing_artifacts(self) -> None:
        with self.assertRaisesRegex(ValueError, "unsupported display kind"):
            parse_displays("chart|project/artifacts/results_v001.csv#T-1|Caption")
        with self.assertRaisesRegex(ValueError, "unsupported file extension"):
            parse_displays("figure|project/artifacts/results_v001.csv#F-1|Caption")
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "source.md"
            source.write_text(sample_source(), encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "do not exist"):
                build(source, root / "out.docx", root / "manifest.json", root)

    def test_rejects_artifact_paths_that_escape_project(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            self._project(root)
            outside = root / "outside.csv"
            outside.write_text("analysis,estimate\nprimary,0.12\n", encoding="utf-8")
            source = root / "source.md"
            source.write_text(
                sample_source("md").replace(
                    "project/artifacts/results_v001.csv#TABLE-PRIMARY",
                    "project/artifacts/../../outside.csv#TABLE-PRIMARY",
                    1,
                ),
                encoding="utf-8",
            )
            with self.assertRaisesRegex(ValueError, "escapes project"):
                build(
                    source,
                    root / "skeleton_draft_v001.md",
                    root / "manifest.json",
                    root,
                )

    def test_refuses_to_overwrite_any_output(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source, output, manifest = self._build(root, "md")
            with self.assertRaisesRegex(ValueError, "refusing to overwrite"):
                build(source, output, manifest, root)

    def test_rejects_unsafe_equation_and_malformed_table(self) -> None:
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            self._project(root)
            equation = root / "project/artifacts/estimating_equation_v001.tex"
            equation.write_text(r"\input{secret.tex}", encoding="utf-8")
            source = root / "source.md"
            source.write_text(sample_source(), encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "unsafe LaTeX command"):
                build(source, root / "out.docx", root / "manifest.json", root)
            equation.write_text(r"Y = X\beta", encoding="utf-8")
            table = root / "project/artifacts/results_v001.csv"
            table.write_text("a,b\n1\n", encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "inconsistent row widths"):
                build(source, root / "out.docx", root / "manifest.json", root)


if __name__ == "__main__":
    unittest.main()
